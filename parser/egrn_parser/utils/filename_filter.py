"""
egrn_parser/utils/filename_filter.py — фильтрация файлов при сканировании.

Правила отфильтровывания DOCX-фотоотчётов (ТЗ раздел 4.1):
  1. Имя файла заканчивается на *-фотоотчёт.docx / *-фотоотчет.docx
  2. Первые 10 параграфов содержат слово «Фотоотчёт»
  3. Количество встроенных изображений / количество таблиц > 5
"""

from __future__ import annotations

import logging
from pathlib import Path

log = logging.getLogger(__name__)

# Суффиксы, признаваемые как фотоотчёт
_PHOTO_SUFFIXES = ("-фотоотчёт.docx", "-фотоотчет.docx", "-фотоотчёт.doc", "-фотоотчет.doc")


def is_photo_report_by_name(path: Path | str) -> bool:
    """Проверить по имени файла, является ли он фотоотчётом."""
    name = Path(path).name.lower()
    return any(name.endswith(suf.lower()) for suf in _PHOTO_SUFFIXES)


def is_photo_report_by_content(path: Path | str) -> bool:
    """
    Проверить по содержимому DOCX, является ли он фотоотчётом.
    Критерии: ключевое слово в первых 10 параграфах ИЛИ соотношение изображений/таблиц > 5.
    Импортирует docx лениво — только при вызове.
    """
    try:
        import docx  # python-docx
        doc = docx.Document(str(path))
    except Exception:
        return False

    # Критерий 1: ключевые слова в начале
    for i, para in enumerate(doc.paragraphs[:10]):
        text_lower = para.text.lower()
        if "фотоотчёт" in text_lower or "фотоотчет" in text_lower or "photo report" in text_lower:
            return True

    # Критерий 2: соотношение изображений/таблиц
    from docx.oxml.ns import qn
    inline_shapes = doc.inline_shapes
    img_count = sum(
        1 for s in inline_shapes
        if s.type.name in ("PICTURE", "LINKED_PICTURE")
    )
    table_count = len(doc.tables)
    if table_count == 0 and img_count > 5:
        return True
    if table_count > 0 and img_count / table_count > 5.0:
        return True

    return False


def is_photo_report(path: Path | str) -> bool:
    """
    Комбинированная проверка: сначала по имени, затем по содержимому.
    """
    path = Path(path)
    if is_photo_report_by_name(path):
        log.info(
            "Файл %s распознан как фотоотчёт (по имени); не используется как источник данных",
            path.name,
        )
        return True
    if path.suffix.lower() in (".docx", ".doc"):
        if is_photo_report_by_content(path):
            log.info(
                "Файл %s распознан как фотоотчёт (по содержимому); не используется как источник данных",
                path.name,
            )
            return True
    return False


def filter_source_files(paths: list[Path]) -> tuple[list[Path], list[Path]]:
    """
    Разделить список файлов на «принятые» и «отфильтрованные» (фотоотчёты и пр.).
    Возвращает (accepted, rejected).
    """
    accepted: list[Path] = []
    rejected: list[Path] = []
    for p in paths:
        if p.suffix.lower() in (".docx", ".doc") and is_photo_report(p):
            rejected.append(p)
        else:
            accepted.append(p)
    return accepted, rejected
