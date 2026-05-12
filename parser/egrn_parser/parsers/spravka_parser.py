"""
egrn_parser/parsers/spravka_parser.py — парсер Справки по юридическим вопросам.

Структура документа ООО ССР:
  Таблица 0: Оформление ЗУ под ОКС
    Колонки: №, Земельный участок (с КН), Срок заключения договора аренды, Статус, Комментарии
  Таблица 1: Объекты ОКС (здания со статусами)
    Колонки: №, Объект капитального строительства, Кадастровый номер, Статус

Результат:
  - lease_intentions: list[dict] — будущие договоры аренды ЗУ (right_category='encumbrance')
  - building_statuses: list[dict] — статусы зданий
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any, Optional

from egrn_parser.parsers._common import (
    CAD_NUMBER_RE,
    normalize_cad_number,
    normalize_whitespace,
    parse_date_ru,
    parse_date_any,
    is_absent,
    clean_value,
)

log = logging.getLogger(__name__)

# Маркеры заголовков таблицы аренды
LEASE_TABLE_HEADERS = frozenset({
    "срок заключения договора аренды",
    "срок аренды",
    "договора аренды",
})

# Маркеры заголовков таблицы объектов
OBJ_TABLE_HEADERS = frozenset({
    "объект капитального строительства",
    "кадастровый номер",
})


def is_spravka_docx(path: Path | str) -> bool:
    """Проверить, является ли DOCX справкой по юридическим вопросам."""
    try:
        import docx
        doc = docx.Document(str(path))
    except Exception:
        return False
    # Ищем маркер в первых 5 параграфах
    for para in doc.paragraphs[:5]:
        t = para.text.strip().lower()
        if "справка" in t or "юридическим вопросам" in t:
            return True
    # Или по наличию таблицы с «срок заключения договора аренды»
    for table in doc.tables:
        if table.rows:
            header_text = " ".join(c.text.strip().lower() for c in table.rows[0].cells)
            if any(kw in header_text for kw in LEASE_TABLE_HEADERS):
                return True
    return False


def _extract_cad_from_cell(cell_text: str) -> Optional[str]:
    """Извлечь полный кадастровый номер из ячейки таблицы (после «КН»)."""
    # Варианты: «КН 90:25:020103:9466», «кадастровый номер 90:25:…»
    # Или просто номер в тексте
    m = re.search(r'(?:КН|кн|кадастровый номер)\s+(' + CAD_NUMBER_RE.pattern + r')', cell_text, re.IGNORECASE)
    if m:
        return m.group(1)
    m = CAD_NUMBER_RE.search(cell_text)
    return m.group(0) if m else None


def _extract_cad_from_parens(cell_text: str) -> Optional[str]:
    """Извлечь кад. номер из «на БУСе 90:25:020102:356»."""
    m = re.search(r'(?:БУСе?|на участке)\s+(' + CAD_NUMBER_RE.pattern + r')', cell_text, re.IGNORECASE)
    if m:
        return m.group(1)
    return None


def parse_spravka_docx(path: Path | str, entity_inn: str | None = None) -> dict:
    """
    Разобрать Справку по юридическим вопросам.

    Возвращает dict:
      lease_intentions: list[dict]   — намерения/статусы договоров аренды ЗУ
      building_statuses: list[dict]  — статусы зданий по ОКС-таблице
      metadata: dict                 — дата справки, организация
      warnings: list[str]
      source_file: str
    """
    path = Path(path)
    result: dict[str, Any] = {
        "lease_intentions": [],
        "building_statuses": [],
        "metadata": {},
        "warnings": [],
        "source_file": path.name,
    }

    try:
        import docx
        doc = docx.Document(str(path))
    except Exception as e:
        result["warnings"].append(f"Ошибка открытия: {e}")
        return result

    # Метаданные: организация и дата из параграфов
    for para in doc.paragraphs[:6]:
        t = para.text.strip()
        if t and "справка" not in t.lower() and "юридическим" not in t.lower():
            if re.search(r'\d{2}\.\d{2}\.\d{4}', t):
                result["metadata"]["date"] = parse_date_ru(t)
            elif t and not result["metadata"].get("organization"):
                result["metadata"]["organization"] = clean_value(t)

    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
        header_text  = " ".join(header_cells)

        # ── Таблица аренды ЗУ ────────────────────────────────────────────────
        if any(kw in header_text for kw in LEASE_TABLE_HEADERS):
            _parse_lease_table(table, result, entity_inn, path.name)

        # ── Таблица ОКС ───────────────────────────────────────────────────────
        elif any(kw in header_text for kw in OBJ_TABLE_HEADERS):
            _parse_oks_table(table, result, path.name)

    log.info(
        "✓ Справка %s: %d аренд ЗУ, %d ОКС",
        path.name, len(result["lease_intentions"]), len(result["building_statuses"]),
    )
    return result


def _parse_lease_table(table, result: dict, entity_inn: Optional[str], source_file: str) -> None:
    """Таблица: № | Земельный участок | Срок заключения договора аренды | Статус | Комментарии."""
    headers = [c.text.strip().lower() for c in table.rows[0].cells]

    # Определить индексы колонок
    def col_idx(keywords: list[str]) -> Optional[int]:
        for kw in keywords:
            for i, h in enumerate(headers):
                if kw in h:
                    return i
        return None

    col_num      = col_idx(["№"])
    col_land     = col_idx(["земельный участок", "участок"])
    col_deadline = col_idx(["срок", "договора аренды"])
    col_status   = col_idx(["статус"])
    col_comment  = col_idx(["комментарии", "примечания"])

    for row in table.rows[1:]:
        cells = [c.text.strip().replace("\xa0", " ") for c in row.cells]
        if not any(cells):
            continue

        # Получить текст земельного участка
        land_text = cells[col_land] if col_land is not None and col_land < len(cells) else ""
        if not land_text:
            continue

        cad_number = _extract_cad_from_cell(land_text) or _extract_cad_from_parens(land_text)
        deadline   = cells[col_deadline] if col_deadline is not None and col_deadline < len(cells) else ""
        status     = cells[col_status]   if col_status   is not None and col_status   < len(cells) else ""
        comment    = cells[col_comment]  if col_comment  is not None and col_comment  < len(cells) else ""
        num        = cells[col_num]      if col_num      is not None and col_num      < len(cells) else ""

        # Попытаться распознать дату дедлайна
        deadline_date = parse_date_ru(deadline) if deadline else None
        # Если дедлайн содержит не дату а описание — сохраняем как текст
        deadline_text = deadline if deadline and not deadline_date else None

        # Вычислить площадь из текста ячейки
        area_m = re.search(r'(?:площадь\s+)?([\d\s]+[.,]\d+)\s*(?:кв\.м|м2)', land_text, re.IGNORECASE)
        area = float(area_m.group(1).replace(" ", "").replace(",", ".")) if area_m else None

        rec: dict[str, Any] = {
            "num":                   normalize_whitespace(num),
            "cad_number":            cad_number,
            "description":           normalize_whitespace(land_text),
            "area":                  area,
            "lease_deadline_date":   deadline_date,
            "lease_deadline_text":   deadline_text or deadline,
            "status":                normalize_whitespace(status),
            "comment":               normalize_whitespace(comment),
            "source_file":           source_file,
            # Для записи в rights
            "right_category":        "encumbrance",
            "right_type":            "Аренда",
            "right_type_code":       "lease",
            "beneficiary_name":      None,   # арендатор — сама организация
            "beneficiary_inn":       entity_inn,
            "valid_until":           deadline_date,
            "lease_term_description":deadline_text or deadline,
        }
        result["lease_intentions"].append(rec)


def _parse_oks_table(table, result: dict, source_file: str) -> None:
    """Таблица: № | Объект ОКС | Кадастровый номер | Статус."""
    headers = [c.text.strip().lower() for c in table.rows[0].cells]

    def col_idx(keywords: list[str]) -> Optional[int]:
        for kw in keywords:
            for i, h in enumerate(headers):
                if kw in h:
                    return i
        return None

    col_name   = col_idx(["объект", "наименование"])
    col_cad    = col_idx(["кадастровый номер", "кн"])
    col_status = col_idx(["статус"])

    for row in table.rows[1:]:
        cells = [c.text.strip().replace("\xa0", " ") for c in row.cells]
        if not any(cells):
            continue

        cad_raw  = cells[col_cad]    if col_cad    is not None and col_cad    < len(cells) else ""
        name     = cells[col_name]   if col_name   is not None and col_name   < len(cells) else ""
        status   = cells[col_status] if col_status is not None and col_status < len(cells) else ""

        cad_number = normalize_cad_number(cad_raw)
        if not cad_number:
            continue

        result["building_statuses"].append({
            "cad_number":  cad_number,
            "name":        normalize_whitespace(name),
            "doc_status":  normalize_whitespace(status),
            "source_file": source_file,
        })

# ─────────────────────────────────────────────────────────────────────────────
#  Парсер Перечня имущества
# ─────────────────────────────────────────────────────────────────────────────

PERECHEN_LAND_HEADERS = frozenset({"кадастровый номер", "срок аренды", "площадь и иные"})
PERECHEN_BLDG_HEADERS = frozenset({"наименование по выписке", "фактическое назначение", "год постройки"})


def is_perechen_docx(path) -> bool:
    """Проверить, является ли DOCX перечнем имущества (для обогащения данных)."""
    try:
        import docx
        doc = docx.Document(str(path))
    except Exception:
        return False
    paras = [p.text.strip().lower() for p in doc.paragraphs[:8] if p.text.strip()]
    combined = " ".join(paras[:4])
    if ("перечень" in combined or
        ("объектам недвижимости" in combined and "справка" in combined) or
        ("объектам недвижимости" in combined and "санаторий" in combined.lower())):
        return True
    # Также проверить если в таблицах есть «год постройки» и «кадастровый номер»
    for table in doc.tables[:2]:
        if table.rows:
            h = " ".join(c.text.strip().lower() for c in table.rows[0].cells)
            if "кадастровый номер" in h and ("год постройки" in h or "срок аренды" in h):
                return True
    return False


def parse_perechen_docx(path, entity_inn=None) -> dict:
    """
    Разобрать Перечень имущества ССР.

    Возвращает dict:
      land_leases:      list[dict]  — ЗУ с реальными сроками аренды
      building_enrichments: list[dict] — обогащающие данные для зданий
      warnings: list[str]
      source_file: str
    """
    import re
    import docx as _docx
    from egrn_parser.parsers._common import parse_number, parse_date_ru, normalize_whitespace
    path_obj = Path(path)
    result = {
        "land_leases":          [],
        "building_enrichments": [],
        "warnings":             [],
        "source_file":          path_obj.name,
    }

    try:
        doc = _docx.Document(str(path_obj))
    except Exception as e:
        result["warnings"].append(str(e))
        return result

    CAD_RE    = re.compile(r"(\d{2}:\d{2}:\d{6,7}:\d+)")
    AREA_RE   = re.compile(r"([\d,\s]+(?:\.\d+)?)\s*(?:кв\.м|м2|квм)", re.IGNORECASE)
    FLOORS_RE = re.compile(r"(\d+)\s*эт", re.IGNORECASE)
    YEAR_RE   = re.compile(r"(?:год|лет)\s+(?:постройки|ввода)[^\d]*(\d{4})", re.IGNORECASE)
    PERIOD_RE = re.compile(r"с\s+(\d{2}\.\d{2}\.\d{4})\s+по\s+(\d{2}\.\d{2}\.\d{4})")

    for table in doc.tables:
        if not table.rows:
            continue
        header_text = " ".join(c.text.strip().lower() for c in table.rows[0].cells)

        # ── Таблица ЗУ (аренда) ──────────────────────────────────────────────
        if "срок аренды" in header_text or "срок заключения" in header_text:
            for row in table.rows[1:]:
                cells = [c.text.strip().replace(chr(160), " ") for c in row.cells]
                full_text = " ".join(cells)
                cads = CAD_RE.findall(full_text)
                if not cads:
                    continue
                period_m = PERIOD_RE.search(full_text)
                area_m   = AREA_RE.search(cells[3] if len(cells) > 3 else "")
                result["land_leases"].append({
                    "cad_number":      cads[0],
                    "description":     normalize_whitespace(cells[2]) if len(cells) > 2 else None,
                    "area":            parse_number(area_m.group(1)) if area_m else None,
                    "valid_from":      parse_date_ru(period_m.group(1)) if period_m else None,
                    "valid_until":     parse_date_ru(period_m.group(2)) if period_m else None,
                    "lease_period_raw":normalize_whitespace(cells[4]) if len(cells) > 4 else None,
                    "right_category":  "encumbrance",
                    "right_type":      "Аренда",
                    "beneficiary_inn": entity_inn,
                    "source_file":     path_obj.name,
                })

        # ── Таблица зданий/сооружений ─────────────────────────────────────────
        elif "наименование по выписке" in header_text or "год постройки" in header_text:
            for row in table.rows[1:]:
                cells = [c.text.strip().replace(chr(160), " ") for c in row.cells]
                full_text = " ".join(cells)
                cads = CAD_RE.findall(full_text)
                if not cads:
                    continue
                # Площадь и этажность из col[3]
                chars_text = cells[3] if len(cells) > 3 else ""
                area_m    = AREA_RE.search(chars_text)
                floors_m  = FLOORS_RE.search(chars_text)
                # Год
                year_col = cells[4] if len(cells) > 4 else ""
                year_m   = YEAR_RE.search(year_col + " " + chars_text)
                year     = int(year_m.group(1)) if year_m else None

                result["building_enrichments"].append({
                    "cad_number":       cads[0],
                    "name_egrn":        normalize_whitespace(cells[1])[:100] if len(cells) > 1 else None,
                    "functional_name":  normalize_whitespace(cells[2])[:100] if len(cells) > 2 else None,
                    "area_from_perechen": parse_number(area_m.group(1)) if area_m else None,
                    "floors_from_perechen": int(floors_m.group(1)) if floors_m else None,
                    "year_built_from_perechen": year,
                    "source_file":      path_obj.name,
                })

    log.info("✓ Перечень %s: %d ЗУ, %d зданий",
             path_obj.name, len(result["land_leases"]), len(result["building_enrichments"]))
    return result
