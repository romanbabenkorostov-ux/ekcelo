#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""ekcelo · 12_make_contract_appendix2_v1 — Приложение №2 (DOCX + MD).

Формирует «Приложение №2 (обязательное) к Договору № …»:
  1. Календарный план (№, Этап, Срок, Результат).
  2. Перечень источников информации (№, Источник, Объём, Срок).
  3. Чеклист услуг со стоимостью (№, Услуга, Ед. изм., Цена за ед.,
     Кол-во, Сумма) + строка ИТОГО.
  4. Подписи сторон.

Список услуг — фиксированный (17 позиций); цены и количества читаются
из services_pricing.json (если есть). Иначе — плейсхолдеры «___» и
ИТОГО = «___».

CLI:
  python pirushin_sosn_rocha_12_make_contract_appendix2_v1.py
"""

from __future__ import annotations

import json
import sys
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

# ───────────────────────────── CLI helpers ──────────────────────────────

BANNER = "═" * 66


class C:
    R = "\033[31m"; G = "\033[32m"; Y = "\033[33m"; B = "\033[1m"
    CY = "\033[36m"; X = "\033[0m"


def cp(t: str = "", c: str = "") -> None:
    print(f"{c}{t}{C.X}" if c else t)


def ask(prompt: str, default: str = "") -> str:
    suffix = f" [{default}]" if default else ""
    s = input(f"{prompt}{suffix}: ").strip()
    return s or default


def ask_path(prompt: str, must_exist: bool = True,
             default: str = "") -> Path:
    while True:
        raw = ask(prompt, default)
        if not raw:
            cp("  ⚠ Путь не задан.", C.Y); continue
        p = Path(raw.strip().strip('"').strip("'")).expanduser()
        if must_exist and not p.exists():
            cp(f"  ⚠ Не найдено: {p}", C.Y); continue
        return p


_RU_MONTHS = [
    "", "января", "февраля", "марта", "апреля", "мая", "июня",
    "июля", "августа", "сентября", "октября", "ноября", "декабря",
]


def ru_date(dt: datetime) -> str:
    return f"{dt.day} {_RU_MONTHS[dt.month]} {dt.year}г."


# ───────────────────────────── данные ───────────────────────────────────

# Календарный план — короткий плейсхолдер на 3 этапа.
DEFAULT_CALENDAR: list[tuple[str, str, str]] = [
    (
        "Этап 1. Сбор и систематизация исходных документов",
        "5 рабочих дней с даты подписания",
        "Структура исходных документов, реестр пробелов",
    ),
    (
        "Этап 2. Выезд на идентификацию + фотофиксация",
        "по согласованию с балансодержателем (1-3 дня)",
        "doc-фотофиксация, привязка объектов к местности",
    ),
    (
        "Этап 3. Финальная сборка и передача",
        "3 рабочих дня после Этапа 2",
        "DOCX/PDF/KMZ-KML, граф связей, ссылка на файлообменник",
    ),
]

# Перечень источников информации — плейсхолдеры.
DEFAULT_SOURCES: list[tuple[str, str, str]] = [
    ("Бухгалтерская справка / ОСВ по сч. 01, 03, 08",
     "по перечню объектов", "за 1 раб. день до Этапа 1"),
    ("Технические паспорта БТИ / тех. описание оборудования",
     "по перечню объектов", "за 1 раб. день до Этапа 1"),
    ("Выписки из ЕГРН (актуальные / архивные)",
     "по перечню объектов недвижимости", "при необходимости (Этап 1-2)"),
    ("Договоры собственности / аренды / залога",
     "сводный пакет", "за 1 раб. день до Этапа 1"),
    ("Контактное лицо балансодержателя для выезда",
     "ФИО, телефон, email", "за 2 раб. дня до Этапа 2"),
    ("Прочие документы (по запросу Исполнителя)",
     "по факту запроса", "в течение 2 раб. дней"),
]

# Список услуг — фиксированный по ТЗ заказчика. Ед. измерения —
# подсказка; по факту корректируется в Word после генерации.
SERVICES: list[tuple[str, str]] = [
    ("Формирование структуры исходных документов с систематизацией",
     "комплект"),
    ("Согласование выезда на идентификацию с представителем "
     "балансодержателя", "выезд"),
    ("Простой при ожидании согласования выезда на идентификацию с "
     "представителем балансодержателя", "час"),
    ("Выезд к местоположению объектов", "выезд"),
    ("Фотофиксация объектов в их локации", "объект"),
    ("Ожидание и приём дополнительных материалов для идентификации",
     "час"),
    ("Идентификация вновь выявленных объектов", "объект"),
    ("Привязка к местности объектов без координат в Росреестре "
     "(с консультацией представителя балансодержателя)", "объект"),
    ("Размещение блока фотографий на файлообменнике", "комплект"),
    ("Предоставление ссылки на скачивание", "ссылка"),
    ("Передача документации в письме", "письмо"),
    ("Формирование doc-фотофиксации", "документ"),
    ("Формирование графа связей", "граф"),
    ("Формирование kmz-kml файла", "файл"),
    ("Блок анализа", "блок"),
    ("Возмещение транспортных расходов", "по факту"),
    ("Возмещение проживания вне точки дислокации специалистов "
     "по идентификации", "сутки"),
]


@dataclass
class ServiceRow:
    """Одна строка чеклиста с ценой."""

    name: str
    unit: str
    price: float | None = None     # за единицу
    qty: float | None = None
    sum: float | None = None       # = price * qty (если оба заданы)

    @classmethod
    def placeholder(cls, name: str, unit: str) -> "ServiceRow":
        return cls(name=name, unit=unit)


def _load_pricing(project: Path | None) -> dict[str, dict[str, Any]]:
    """Читает services_pricing.json (если есть). Возвращает map
    name → {price, qty}. Если файл не найден или битый — пустой словарь.
    """
    here = Path(__file__).resolve().parent
    repo = here.parent.parent
    candidates: list[Path] = []
    if project:
        candidates.append(project / "_data" / "services_pricing.json")
    candidates.extend([
        Path.cwd() / "services_pricing.json",
        repo / "parser" / "_data" / "services_pricing.json",
        repo / "obsidian" / "Database" / "services_pricing.json",
    ])
    for p in candidates:
        if not p.is_file():
            continue
        try:
            data = json.loads(p.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError) as e:
            cp(f"  ⚠ Не удалось прочитать {p}: {e}", C.Y); continue
        cp(f"· services_pricing.json: {p}", C.CY)
        return {str(row.get("name", "")): row for row in data.get("rows", [])}
    cp("· services_pricing.json не найден — цены оставлены пустыми.", C.Y)
    return {}


def _build_service_rows(pricing: dict[str, dict[str, Any]]) -> list[ServiceRow]:
    rows: list[ServiceRow] = []
    for name, unit in SERVICES:
        cfg = pricing.get(name) or {}
        price = cfg.get("price")
        qty = cfg.get("qty")
        # Защита: только число / None.
        try:
            price_v = float(price) if price is not None else None
        except (TypeError, ValueError):
            price_v = None
        try:
            qty_v = float(qty) if qty is not None else None
        except (TypeError, ValueError):
            qty_v = None
        sum_v: float | None = None
        if price_v is not None and qty_v is not None:
            sum_v = round(price_v * qty_v, 2)
        rows.append(ServiceRow(name=name, unit=unit, price=price_v,
                               qty=qty_v, sum=sum_v))
    return rows


def _fmt_money(v: float | None) -> str:
    if v is None:
        return "___"
    # 12345.6 → «12 345,60»
    int_part = int(abs(v))
    int_str = f"{int_part:,}".replace(",", " ")
    sign = "-" if v < 0 else ""
    frac = abs(v) - int_part
    return f"{sign}{int_str},{int(round(frac*100)):02d}"


def _fmt_qty(v: float | None) -> str:
    if v is None:
        return "___"
    if abs(v - round(v)) < 1e-9:
        return str(int(round(v)))
    return f"{v:.2f}".replace(".", ",")


# ───────────────────────────── DOCX builder ─────────────────────────────

def build_docx(out_path: Path, contract_num: str, contract_date: str,
               services: list[ServiceRow]) -> Path:
    from docx import Document  # type: ignore[import]
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for sec in doc.sections:
        sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
        sec.top_margin = sec.bottom_margin = Cm(2.0)
        sec.left_margin = sec.right_margin = Cm(2.0)
    st = doc.styles["Normal"]
    st.font.name = "Arial"; st.font.size = Pt(10)

    # Шапка: «Приложение №2 (обязательное)» курсивом по правому краю +
    # «к Договору № …» — следующей строкой того же параграфа.
    appendix = doc.add_paragraph()
    appendix.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = appendix.add_run("Приложение №2 (обязательное)")
    r1.italic = True
    r1.add_break()
    r2 = appendix.add_run(f"к Договору № {contract_num} от {contract_date}")
    r2.italic = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(
        "Календарный план, Перечень источников информации, "
        "Чеклист услуг со стоимостью"
    )
    tr.bold = True

    # — 1. Календарный план
    _emit_calendar_table(doc, DEFAULT_CALENDAR)
    # — 2. Перечень источников
    _emit_sources_table(doc, DEFAULT_SOURCES)
    # — 3. Чеклист услуг + ИТОГО
    _emit_services_table(doc, services)

    # — Подписи
    doc.add_paragraph()
    _emit_signatures_block(doc)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def _add_section_heading(doc, text: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import]
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text); r.bold = True


def _emit_calendar_table(doc, rows: list[tuple[str, str, str]]) -> None:
    _add_section_heading(doc, "Таблица 1. Календарный план")
    headers = ("№", "Этап", "Срок", "Результат")
    widths = (1.0, 7.0, 4.5, 4.5)  # 17 см
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"; tbl.autofit = False
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(h); run.bold = True
    _mark_row_as_header(tbl.rows[0])
    for i, (stage, term, result) in enumerate(rows, 1):
        row = tbl.add_row()
        row.cells[0].text = str(i)
        row.cells[1].text = stage
        row.cells[2].text = term
        row.cells[3].text = result
    _set_column_widths(tbl, widths)


def _emit_sources_table(doc, rows: list[tuple[str, str, str]]) -> None:
    _add_section_heading(doc, "Таблица 2. Перечень источников информации")
    headers = ("№", "Источник", "Объём", "Срок предоставления")
    widths = (1.0, 8.0, 3.5, 4.5)
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"; tbl.autofit = False
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(h); run.bold = True
    _mark_row_as_header(tbl.rows[0])
    for i, (src, vol, term) in enumerate(rows, 1):
        row = tbl.add_row()
        row.cells[0].text = str(i)
        row.cells[1].text = src
        row.cells[2].text = vol
        row.cells[3].text = term
    _set_column_widths(tbl, widths)


def _emit_services_table(doc, rows: list[ServiceRow]) -> None:
    _add_section_heading(
        doc, "Таблица 3. Чеклист услуг (стоимость единицы × количество)"
    )
    headers = ("№", "Услуга", "Ед. изм.", "Цена за ед., руб.",
               "Кол-во", "Сумма, руб.")
    widths = (0.8, 7.0, 1.6, 2.8, 1.5, 3.3)  # 17 см
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"; tbl.autofit = False
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(h); run.bold = True
    _mark_row_as_header(tbl.rows[0])

    total = 0.0
    has_any_sum = False
    for i, r in enumerate(rows, 1):
        row = tbl.add_row()
        row.cells[0].text = str(i)
        row.cells[1].text = r.name
        row.cells[2].text = r.unit
        row.cells[3].text = _fmt_money(r.price)
        row.cells[4].text = _fmt_qty(r.qty)
        row.cells[5].text = _fmt_money(r.sum)
        if r.sum is not None:
            total += r.sum; has_any_sum = True

    # Строка ИТОГО — выделена жирным, последняя ячейка содержит сумму.
    total_row = tbl.add_row()
    total_row.cells[0].text = ""
    merged = total_row.cells[0]
    # Сливаем первые 5 ячеек под подпись «ИТОГО».
    for j in range(1, 5):
        merged = merged.merge(total_row.cells[j])
    merged.text = ""
    merged.paragraphs[0].add_run("ИТОГО:").bold = True
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import]
    merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    total_row.cells[5].text = ""
    total_run = total_row.cells[5].paragraphs[0].add_run(
        _fmt_money(total) if has_any_sum else "___"
    )
    total_run.bold = True
    _set_cell_shading(total_row.cells[5], "F2F2F2")
    _set_cell_shading(merged, "F2F2F2")

    _set_column_widths(tbl, widths)


def _emit_signatures_block(doc) -> None:
    doc.add_paragraph("Подписи сторон:")
    doc.add_paragraph()
    tbl = doc.add_table(rows=2, cols=2)
    tbl.autofit = False
    tbl.cell(0, 0).text = "От Исполнителя:"
    tbl.cell(0, 1).text = "От Заказчика:"
    tbl.cell(1, 0).text = "_________/_________________/"
    tbl.cell(1, 1).text = "_________/_________________/"
    _set_column_widths(tbl, (8.5, 8.5))
    _remove_table_borders(tbl)


# ─────────────────────── DOCX low-level helpers ─────────────────────────

def _mark_row_as_header(row) -> None:
    """tblHeader → строка повторяется на каждой странице."""
    from docx.oxml.ns import qn  # type: ignore[import]
    from docx.oxml import OxmlElement  # type: ignore[import]
    tr = row._tr
    tr_pr = tr.find(qn("w:trPr"))
    if tr_pr is None:
        tr_pr = OxmlElement("w:trPr"); tr.insert(0, tr_pr)
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def _set_cell_shading(cell, fill_hex: str) -> None:
    from docx.oxml.ns import qn  # type: ignore[import]
    from docx.oxml import OxmlElement  # type: ignore[import]
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _remove_table_borders(tbl) -> None:
    from docx.oxml.ns import qn  # type: ignore[import]
    from docx.oxml import OxmlElement  # type: ignore[import]
    tbl_el = tbl._tbl
    tbl_pr = tbl_el.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr"); tbl_el.insert(0, tbl_pr)
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tbl_pr.append(borders)


def _set_column_widths(tbl, widths_cm: tuple[float, ...]) -> None:
    from docx.shared import Cm  # type: ignore[import]
    for col, w in enumerate(widths_cm):
        for cell in tbl.column_cells(col):
            cell.width = Cm(w)


# ───────────────────────────── MD builder ───────────────────────────────

def build_md(out_path: Path, contract_num: str, contract_date: str,
             services: list[ServiceRow]) -> Path:
    lines: list[str] = []
    lines.append(
        f"*Приложение №2 (обязательное) "
        f"к Договору № {contract_num} от {contract_date}*"
    )
    lines.append("")
    lines.append(
        "# Календарный план, Перечень источников информации, "
        "Чеклист услуг со стоимостью"
    )
    lines.append("")

    lines.append("## Таблица 1. Календарный план")
    lines.append("")
    lines.append("| № | Этап | Срок | Результат |")
    lines.append("|---|---|---|---|")
    for i, (stage, term, result) in enumerate(DEFAULT_CALENDAR, 1):
        lines.append(f"| {i} | {stage} | {term} | {result} |")
    lines.append("")

    lines.append("## Таблица 2. Перечень источников информации")
    lines.append("")
    lines.append("| № | Источник | Объём | Срок |")
    lines.append("|---|---|---|---|")
    for i, (src, vol, term) in enumerate(DEFAULT_SOURCES, 1):
        lines.append(f"| {i} | {src} | {vol} | {term} |")
    lines.append("")

    lines.append("## Таблица 3. Чеклист услуг")
    lines.append("")
    lines.append(
        "| № | Услуга | Ед. изм. | Цена за ед., руб. | "
        "Кол-во | Сумма, руб. |"
    )
    lines.append("|---|---|---|---:|---:|---:|")
    total = 0.0; has_any = False
    for i, r in enumerate(services, 1):
        lines.append(
            f"| {i} | {r.name} | {r.unit} | "
            f"{_fmt_money(r.price)} | {_fmt_qty(r.qty)} | "
            f"{_fmt_money(r.sum)} |"
        )
        if r.sum is not None:
            total += r.sum; has_any = True
    lines.append(
        f"| | **ИТОГО:** | | | | "
        f"**{_fmt_money(total) if has_any else '___'}** |"
    )
    lines.append("")

    lines.append("**Подписи сторон:**")
    lines.append("")
    lines.append("| От Исполнителя | От Заказчика |")
    lines.append("|---|---|")
    lines.append("| _________/_________/ | _________/_________/ |")
    lines.append("")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(lines), encoding="utf-8")
    return out_path


# ───────────────────────────── main ─────────────────────────────────────

def main() -> int:
    cp(f"\n{BANNER}", C.B)
    cp("  ekcelo · Приложение №2 (DOCX + MD)", C.B)
    cp(f"{BANNER}\n", C.B)

    out_dir = ask_path(
        "Папка для сохранения DOCX/MD",
        must_exist=False, default=str(Path.cwd() / "out"),
    )
    out_dir.mkdir(parents=True, exist_ok=True)

    proj_raw = ask("Папка проекта (Enter — пропустить)", "")
    project = Path(proj_raw).expanduser() if proj_raw else None
    if project and not project.is_dir():
        cp(f"  ⚠ Папка не найдена: {project} — игнорирую.", C.Y)
        project = None

    now = datetime.now()
    default_num = f"{now.strftime('%Y%m%d')}-1"
    contract_num = ask("Номер договора", default=default_num)
    contract_date = ask("Дата договора", default=ru_date(now))

    pricing = _load_pricing(project)
    services = _build_service_rows(pricing)

    ts = now.strftime("%Y%m%d_%H%M%S")
    safe_num = contract_num.replace("/", "-").replace(" ", "_")
    out_md = out_dir / f"Appendix2_{safe_num}_{ts}.md"
    out_docx = out_dir / f"Appendix2_{safe_num}_{ts}.docx"

    build_md(out_md, contract_num, contract_date, services)
    cp(f"\n✓ MD : {out_md}", C.G)

    try:
        build_docx(out_docx, contract_num, contract_date, services)
        cp(f"✓ DOCX: {out_docx}", C.G)
    except ImportError as e:
        cp(f"\n⚠ python-docx не установлен: {e}", C.Y)
        cp("  pip install python-docx", C.Y)
        cp(f"  DOCX пропущен; MD доступен: {out_md}", C.Y)
    return 0


if __name__ == "__main__":
    try:
        sys.exit(main())
    except KeyboardInterrupt:
        cp("\n✗ Прервано пользователем.", C.Y)
        sys.exit(130)
