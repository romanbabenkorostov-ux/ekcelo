#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для рекурсивного обхода папки с фотографиями и формирования
docx-фотоотчёта.

  • спрашивает путь к корневой папке (с подтверждением);
  • рекурсивно обходит ВСЕ вложенные папки;
  • для каждой папки, содержащей фотографии, выводит цепочку заголовков:
        Заголовок 2 — имя корневой папки,
        Заголовок 3 — первая вложенность, и т. д. (до Заголовка 9);
  • кладёт фотографии по 3 на лист, A4, поля 2 см;
  • каждое фото вписывается по ширине/высоте слота с сохранением пропорций;
  • маленькие фото не растягиваются (остаются в натуральный размер @ 96 dpi);
  • EXIF-ориентация применяется автоматически;
  • под каждым фото — подпись:
        «<заголовки через тире> - <имя файла>»
        + (если есть EXIF): дата/время съёмки и GPS-координаты;
  • итоговый файл — "<имя_корневой_папки>-фотоотчёт.docx" в самой
    корневой папке.

Совместимо: Python 3.13+, Windows 10, VSCode terminal, MS Word.
Зависимости (уже стоят): python-docx, pillow.
"""

from __future__ import annotations

import sys
import tempfile
import shutil
import warnings
from pathlib import Path
from datetime import datetime

from PIL import Image, ImageOps
from PIL.ExifTags import TAGS, GPSTAGS

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings('ignore')


# ═══════════════════════════════════════════════════════════════════════════
#  Цветной вывод
# ═══════════════════════════════════════════════════════════════════════════
class C:
    GREEN = '\033[92m'
    RED = '\033[91m'
    YELLOW = '\033[93m'
    CYAN = '\033[96m'
    BOLD = '\033[1m'
    RESET = '\033[0m'


def cp(text: str = '', color: str = C.RESET) -> None:
    print(f"{color}{text}{C.RESET}")


# ═══════════════════════════════════════════════════════════════════════════
#  Константы
# ═══════════════════════════════════════════════════════════════════════════
IMAGE_EXTS = {'.jpg', '.jpeg', '.png', '.tif', '.tiff', '.bmp', '.webp'}

# Папки, которые надо игнорировать в подписи к фото
# (в дереве заголовков и TOC они остаются — игнорируется только текст подписи).
IGNORE_FOLDERS = {'фото', 'фотографии', 'фотография', 'foto'}

# A4: 21.0 × 29.7 см. Поля по 2 см → область контента 17 × 25.7 см.
PAGE_MARGIN_CM = 2.0
CONTENT_W_CM = 21.0 - 2 * PAGE_MARGIN_CM       # 17.0
CONTENT_H_CM = 29.7 - 2 * PAGE_MARGIN_CM       # 25.7

# 3 фото на страницу. Слот = ~ 1/3 высоты, минус подпись и небольшой
# резерв на возможный заголовок в начале страницы.
PHOTOS_PER_PAGE = 3
CAPTION_RESERVE_CM = 1.5    # подпись (1–2 строки) + межстрочный отступ
HEADING_RESERVE_CM = 1.0    # запас под цепочку заголовков на стыке секций
SLOT_W_CM = CONTENT_W_CM
SLOT_H_CM = (CONTENT_H_CM - HEADING_RESERVE_CM) / PHOTOS_PER_PAGE - CAPTION_RESERVE_CM
# = (25.7 - 1.0) / 3 - 1.5 ≈ 6.73 см

# Натуральный размер изображения вычисляется как pixels / 96 dpi → см.
ASSUMED_DPI = 96.0


# ═══════════════════════════════════════════════════════════════════════════
#  EXIF
# ═══════════════════════════════════════════════════════════════════════════
def _gps_to_decimal(dms, ref) -> float | None:
    if dms is None or ref is None or len(dms) < 3:
        return None
    try:
        d, m, s = float(dms[0]), float(dms[1]), float(dms[2])
        dec = d + m / 60.0 + s / 3600.0
        if isinstance(ref, bytes):
            ref = ref.decode(errors='ignore')
        if ref in ('S', 'W'):
            dec = -dec
        return dec
    except (TypeError, ValueError):
        return None


def extract_exif(img_path: Path) -> dict:
    """Возвращает {'datetime': str|None, 'gps': (lat, lon)|None, 'orientation': int}."""
    out = {'datetime': None, 'gps': None, 'orientation': 1}
    try:
        with Image.open(img_path) as img:
            exif = img.getexif()
            if not exif:
                return out

            out['orientation'] = exif.get(0x0112) or 1

            # DateTimeOriginal приоритетней, чем просто DateTime
            dt_str = None
            try:
                exif_ifd = exif.get_ifd(0x8769)
                for tid, val in (exif_ifd or {}).items():
                    if TAGS.get(tid) == 'DateTimeOriginal':
                        dt_str = val
                        break
            except Exception:
                pass
            if not dt_str:
                for tid, val in exif.items():
                    if TAGS.get(tid) == 'DateTime':
                        dt_str = val
                        break

            if dt_str:
                if isinstance(dt_str, bytes):
                    dt_str = dt_str.decode(errors='ignore')
                dt_str = dt_str.strip().rstrip('\x00')
                try:
                    dt = datetime.strptime(dt_str, '%Y:%m:%d %H:%M:%S')
                    out['datetime'] = dt.strftime('%d.%m.%Y %H:%M:%S')
                except (ValueError, TypeError):
                    out['datetime'] = dt_str or None

            try:
                gps_ifd = exif.get_ifd(0x8825)
                if gps_ifd:
                    gps = {GPSTAGS.get(k, k): v for k, v in gps_ifd.items()}
                    lat = _gps_to_decimal(gps.get('GPSLatitude'), gps.get('GPSLatitudeRef'))
                    lon = _gps_to_decimal(gps.get('GPSLongitude'), gps.get('GPSLongitudeRef'))
                    if lat is not None and lon is not None:
                        out['gps'] = (lat, lon)
            except Exception:
                pass
    except Exception:
        pass
    return out


# ═══════════════════════════════════════════════════════════════════════════
#  Работа с изображениями
# ═══════════════════════════════════════════════════════════════════════════
def normalize_for_docx(img_path: Path, tmp_dir: Path) -> Path:
    """
    Применяет EXIF-ориентацию (Word её не всегда учитывает).
    Возвращает путь к нормализованной копии (или исходник, если ротация не нужна).
    """
    try:
        with Image.open(img_path) as img:
            exif = img.getexif()
            orient = exif.get(0x0112) if exif else None
            if orient in (3, 6, 8):
                rotated = ImageOps.exif_transpose(img)
                out = tmp_dir / f"{img_path.stem}__rot.jpg"
                if rotated.mode in ('RGBA', 'P', 'LA'):
                    rotated = rotated.convert('RGB')
                rotated.save(out, 'JPEG', quality=92)
                return out
    except Exception:
        pass
    return img_path


def computed_size_cm(img_path: Path) -> tuple[float, float]:
    """
    Финальный размер картинки на странице (см), с учётом:
      - подгонки под слот SLOT_W_CM × SLOT_H_CM по обоим измерениям;
      - запрета увеличения ниже натурального размера @ ASSUMED_DPI.
    Учитывает EXIF-ориентацию (если фото повёрнуто, меняем местами w/h).
    """
    try:
        with Image.open(img_path) as img:
            w_px, h_px = img.size
            exif = img.getexif()
            orient = exif.get(0x0112) if exif else None
            if orient in (5, 6, 7, 8):
                w_px, h_px = h_px, w_px
    except Exception:
        # Если что-то пошло не так — пусть Word отрисует «как есть» в максимум слота.
        return SLOT_W_CM, SLOT_H_CM

    w_native = w_px / ASSUMED_DPI * 2.54
    h_native = h_px / ASSUMED_DPI * 2.54

    scale = min(SLOT_W_CM / w_native, SLOT_H_CM / h_native, 1.0)
    return w_native * scale, h_native * scale


# ═══════════════════════════════════════════════════════════════════════════
#  Сбор данных по дереву папок
# ═══════════════════════════════════════════════════════════════════════════
def list_images(folder: Path) -> list[Path]:
    """Файлы-изображения непосредственно в folder, отсортированные по имени."""
    try:
        return sorted(
            (f for f in folder.iterdir()
             if f.is_file() and f.suffix.lower() in IMAGE_EXTS),
            key=lambda p: p.name.lower(),
        )
    except (PermissionError, OSError):
        return []


def list_subfolders(folder: Path) -> list[Path]:
    try:
        return sorted(
            (f for f in folder.iterdir() if f.is_dir()),
            key=lambda p: p.name.lower(),
        )
    except (PermissionError, OSError):
        return []


def has_descendant_images(folder: Path) -> bool:
    """Есть ли хоть одно фото в этой папке или внутри неё (рекурсивно)?"""
    if list_images(folder):
        return True
    for sub in list_subfolders(folder):
        if has_descendant_images(sub):
            return True
    return False


def count_all_images(folder: Path) -> int:
    n = len(list_images(folder))
    for sub in list_subfolders(folder):
        n += count_all_images(sub)
    return n


# ═══════════════════════════════════════════════════════════════════════════
#  Построение docx
# ═══════════════════════════════════════════════════════════════════════════
def setup_document() -> Document:
    """Создаёт документ с полями 2 см на A4."""
    doc = Document()
    for section in doc.sections:
        # A4
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        # Поля по 2 см
        section.top_margin = Cm(PAGE_MARGIN_CM)
        section.bottom_margin = Cm(PAGE_MARGIN_CM)
        section.left_margin = Cm(PAGE_MARGIN_CM)
        section.right_margin = Cm(PAGE_MARGIN_CM)
        section.header_distance = Cm(0.5)
        section.footer_distance = Cm(0.5)

    # Чуть подкомпактим стиль Normal, чтобы 3 фото точно влезали
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    return doc


def _format_folder_name(name: str) -> str:
    """Универсальная нормализация имени папки для отображения:
    подчёркивания → пробелы, первая буква — заглавной."""
    s = name.replace('_', ' ').strip()
    if s and s[0].isalpha() and s[0].islower():
        s = s[0].upper() + s[1:]
    return s


def build_caption_text(
    chain: list[str],
    filename: str,
    dt: str | None,
    gps: tuple[float, float] | None,
) -> str:
    """
    Собирает текст подписи БЕЗ префикса «Рисунок N.» (он добавляется отдельно
    через SEQ-поле). Цепочка `chain` уже отфильтрована (без корневой папки и
    без IGNORE_FOLDERS) и отформатирована (см. collect_sections), поэтому
    здесь только склейка.

    Формат:
      <Папка1, Папка2, …,> файл <имя>[ (съемка от <дата>, координаты ~ <lat>, <lon>)].
    """
    folders = chain[1:]  # без корневой
    parts = list(folders)
    parts.append(f'файл {filename}')
    main = ', '.join(parts)

    extras: list[str] = []
    if dt:
        extras.append(f'съемка от {dt}')
    if gps:
        lat, lon = gps
        extras.append(f'координаты ~ {lat:.6f}, {lon:.6f}')

    if extras:
        main += ' (' + ', '.join(extras) + ')'

    return main + '.'


def _insert_seq_field(paragraph, seq_name: str, cached_value: str = '#') -> None:
    """
    Вставляет в параграф runs с полем SEQ для автонумерации.
    `cached_value` — что будет показано до обновления полей (F9 / ПКМ → «Обновить»).
    """
    # begin
    r = paragraph.add_run()
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')
    r._element.append(fld)

    # instrText: «SEQ Рисунок \* ARABIC»
    r = paragraph.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' SEQ {seq_name} \\* ARABIC '
    r._element.append(instr)

    # separate
    r = paragraph.add_run()
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'separate')
    r._element.append(fld)

    # cached/placeholder value
    paragraph.add_run(cached_value)

    # end
    r = paragraph.add_run()
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'end')
    r._element.append(fld)


def add_image_caption(doc: Document, caption_text: str, fig_num: int) -> None:
    """
    Добавляет подпись «Рисунок N. <text>» под рисунком.
    Номер N — автонумеруемое поле SEQ; до F9/обновления показывает fig_num.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)

    def _styled(run):
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        return run

    # «Рисунок » с неразрывным пробелом перед номером
    _styled(p.add_run('Рисунок\u00a0'))

    # SEQ-поле, в качестве placeholder — реальный текущий номер
    _insert_seq_field(p, 'Рисунок', cached_value=str(fig_num))
    # Стилизуем placeholder-run (он добавлен последним до end)
    for run in p.runs:
        _styled(run)

    # «. <текст>»
    _styled(p.add_run(f'. {caption_text}'))


def add_photo_block(
    doc: Document,
    src_path: Path,
    docx_img_path: Path,
    header_chain: list[str],
    fig_num: int,
) -> None:
    """Картинка по центру + подпись «Рисунок N. …» под ней."""
    w_cm, h_cm = computed_size_cm(src_path)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    try:
        run.add_picture(str(docx_img_path), width=Cm(w_cm), height=Cm(h_cm))
    except Exception as e:
        cp(f"  ⚠ Не удалось вставить {src_path.name}: {e}", C.YELLOW)
        return

    exif = extract_exif(src_path)
    caption = build_caption_text(
        chain=header_chain,
        filename=src_path.name,
        dt=exif['datetime'],
        gps=exif['gps'],
    )
    add_image_caption(doc, caption, fig_num=fig_num)


def add_heading_safe(doc: Document, text: str, level: int) -> None:
    """add_heading с защитой от слишком глубокой вложенности (cap 9)."""
    level = max(1, min(level, 9))
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def insert_page_break(doc: Document) -> None:
    """Жёсткий разрыв страницы."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_table_of_contents(doc: Document) -> None:
    """
    Добавляет в начало документа заголовок «Оглавление» (Heading 1, чтобы он
    сам не попал в TOC) + поле TOC с уровнями 2..9. Затем — разрыв страницы.
    Для заполнения оглавления достаточно открыть файл в Word и нажать F9
    (или ПКМ по полю → «Обновить поле»).
    """
    h = doc.add_heading('Оглавление', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # begin (с dirty=true — Word поймёт, что поле требует обновления)
    r = p.add_run()
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')
    fld.set(qn('w:dirty'), 'true')
    r._element.append(fld)

    # instrText: TOC \o "2-9" \h \z \u
    r = p.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = r' TOC \o "2-9" \h \z \u '
    r._element.append(instr)

    # separate
    r = p.add_run()
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'separate')
    r._element.append(fld)

    # placeholder (виден до первого обновления полей)
    placeholder = p.add_run(
        'Оглавление будет заполнено после обновления полей '
        '(F9 или ПКМ → «Обновить поле»).'
    )
    placeholder.italic = True
    placeholder.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # end
    r = p.add_run()
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'end')
    r._element.append(fld)

    # Разрыв страницы — чтобы первая фото-секция начиналась с нового листа
    pb = doc.add_paragraph()
    pb.add_run().add_break(WD_BREAK.PAGE)


# ═══════════════════════════════════════════════════════════════════════════
#  Главный обход: сначала собираем плоский список «секций», затем эмитим
# ═══════════════════════════════════════════════════════════════════════════
class BuildState:
    def __init__(self) -> None:
        self.emitted: list[str] = []      # последняя выведенная цепочка заголовков
        self.images_on_page: int = 0      # сколько фото уже на текущей странице
        self.any_images_emitted: bool = False
        self.total_processed: int = 0
        self.fig_counter: int = 0         # placeholder-номер для SEQ-поля


def collect_sections(folder: Path, root: Path,
                     out: list[tuple[list[str], list[Path]]]) -> None:
    """
    DFS-обход. Для каждой папки, у которой ЕСТЬ прямые изображения,
    добавляет в `out` пару (chain_of_headings, list_of_images), где chain:
      • начинается с корневой папки;
      • из подпути вырезаны папки из IGNORE_FOLDERS (Фото/Фотографии/Foto);
      • все элементы нормализованы: «_» → пробел, первая буква заглавной.

    Сама структура файловой системы при этом обходится полностью —
    исключения работают только на уровне отображения.
    """
    direct = list_images(folder)
    if direct:
        if folder == root:
            raw = [root.name]
        else:
            sub = [p for p in folder.relative_to(root).parts
                   if p.lower() not in IGNORE_FOLDERS]
            raw = [root.name] + sub
        chain = [_format_folder_name(p) for p in raw]
        out.append((chain, direct))

    for sub in list_subfolders(folder):
        collect_sections(sub, root, out)


def emit_sections(
    doc: Document,
    sections: list[tuple[list[str], list[Path]]],
    root: Path,
    state: BuildState,
    tmp_dir: Path,
    total_count: int,
) -> None:
    """
    Эмитит секции в документ. Для каждой секции:
      • если уже размещали фото — разрыв страницы перед первым новым заголовком;
      • заголовки, которых ещё не было, выводятся подряд (Heading 2/3/4/…);
      • затем 3 фото на страницу.
    """
    for chain, images in sections:
        # Какие из элементов цепочки уже эмитированы?
        new_start = 0
        for i in range(min(len(chain), len(state.emitted))):
            if chain[i] != state.emitted[i]:
                break
            new_start = i + 1
        new_headings = chain[new_start:]

        # Разрыв страницы перед НОВЫМИ заголовками этой секции,
        # если мы уже что-то размещали (фото).
        if new_headings and state.any_images_emitted:
            insert_page_break(doc)
            state.images_on_page = 0

        for i, name in enumerate(new_headings):
            level = new_start + i + 2   # 2 — для корня
            if level == 2:
                heading_text = f'Фотофиксация при идентификации объектов {name}'
            else:
                heading_text = name
            add_heading_safe(doc, heading_text, level)
        state.emitted = list(chain)

        # Прямые изображения этой папки — по 3 на страницу
        for img in images:
            if state.images_on_page >= PHOTOS_PER_PAGE:
                insert_page_break(doc)
                state.images_on_page = 0

            norm = normalize_for_docx(img, tmp_dir)
            state.fig_counter += 1
            add_photo_block(doc, img, norm, chain, fig_num=state.fig_counter)
            state.images_on_page += 1
            state.any_images_emitted = True
            state.total_processed += 1
            cp(f"  [{state.total_processed}/{total_count}] "
               f"{img.relative_to(root)}", C.GREEN)


def update_fields_via_word(docx_path: Path) -> bool:
    """
    Открывает .docx в MS Word через COM и обновляет ВСЕ поля
    (оглавление + автонумерацию SEQ). После этого в Word файл уже
    отрисуется корректно без F9.

    Возвращает True, если обновление выполнено. Если Word недоступен
    (не Windows / не установлен / занят и т.п.) — печатает предупреждение
    и возвращает False; файл остаётся валидным, поля обновятся при F9.
    """
    try:
        import win32com.client                              # type: ignore[import]
        import pythoncom                                    # type: ignore[import]
    except ImportError:
        cp("  ⚠ pywin32 не найден — оглавление и номера рисунков обновятся при F9.",
           C.YELLOW)
        return False

    word = None
    doc = None
    try:
        pythoncom.CoInitialize()
        try:
            word = win32com.client.DispatchEx('Word.Application')
        except Exception as e:
            cp(f"  ⚠ Не удалось запустить MS Word: {e}", C.YELLOW)
            cp("    Откройте файл в Word и нажмите Ctrl+A → F9.", C.YELLOW)
            return False

        word.Visible = False
        word.DisplayAlerts = False

        # Word требует абсолютного пути и понимает обратные слэши
        doc = word.Documents.Open(str(docx_path.resolve()), ReadOnly=False)

        # Обновить SEQ-поля и т.п.
        try:
            doc.Fields.Update()
        except Exception:
            pass

        # Обновить все оглавления (TOC) — двумя проходами,
        # т.к. вторая итерация устаканивает номера страниц
        try:
            for _ in range(2):
                for toc in doc.TablesOfContents:
                    toc.Update()
        except Exception:
            pass

        doc.Save()
        return True

    except Exception as e:
        cp(f"  ⚠ Ошибка при обновлении полей через Word: {e}", C.YELLOW)
        return False

    finally:
        try:
            if doc is not None:
                doc.Close(SaveChanges=False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


# ═══════════════════════════════════════════════════════════════════════════
#  Консольный интерфейс
# ═══════════════════════════════════════════════════════════════════════════
def ask_yn(question: str, default: bool = True) -> bool:
    suffix = ' (Y/n): ' if default else ' (y/N): '
    while True:
        ans = input(question + suffix).strip().lower()
        if not ans:
            return default
        if ans in ('y', 'yes', 'д', 'да'):
            return True
        if ans in ('n', 'no', 'н', 'нет'):
            return False
        cp("Введите 'y' или 'n'", C.YELLOW)


def ask_dir(prompt: str) -> Path | None:
    while True:
        raw = input(prompt).strip().strip('"').strip("'")
        if not raw:
            cp("Путь не может быть пустым.", C.RED)
            if not ask_yn("Попробовать ещё раз?", default=True):
                return None
            continue
        p = Path(raw)
        if not p.exists():
            cp(f"Папка не существует: {p}", C.RED)
            if not ask_yn("Попробовать ещё раз?", default=True):
                return None
            continue
        if not p.is_dir():
            cp(f"Это не папка: {p}", C.RED)
            if not ask_yn("Попробовать ещё раз?", default=True):
                return None
            continue
        return p.resolve()


def main() -> None:
    cp("\n" + "═" * 70, C.BOLD)
    cp("  ФОТООТЧЁТ ИЗ ПАПКИ → DOCX", C.BOLD)
    cp("═" * 70 + "\n", C.BOLD)

    cp("ШАГ 1. Укажите корневую папку с фотографиями.", C.CYAN)
    cp('Пример: D:\\ОБЪЕКТЫ\\Русаков\\2026-04-28_Суворова\\17_объектов_РнД', C.CYAN)
    root = ask_dir("Путь к папке: ")
    if not root:
        cp("Операция отменена.", C.RED)
        return

    out_path = root / f"{root.name}-фотоотчёт.docx"

    # Подсчёт фото
    cp(f"\nКорневая папка: {root}", C.CYAN)
    cp("Подсчёт фотографий…", C.CYAN)
    total = count_all_images(root)
    cp(f"Найдено фото:   {total}", C.GREEN if total else C.RED)
    cp(f"Будет создан:   {out_path.name}", C.CYAN)

    if total == 0:
        cp("\nФотографий не найдено. Завершаю работу.", C.YELLOW)
        return

    if out_path.exists():
        cp(f"⚠ Файл {out_path.name} уже существует и будет перезаписан.", C.YELLOW)

    if not ask_yn("\nПродолжить?", default=True):
        cp("Операция отменена.", C.RED)
        return

    # Поехали
    cp(f"\n{'─' * 70}", C.BOLD)
    cp("ШАГ 2. Сборка документа", C.BOLD)
    cp("─" * 70, C.BOLD)

    doc = setup_document()
    add_table_of_contents(doc)
    state = BuildState()
    tmp_dir = Path(tempfile.mkdtemp(prefix='photoreport_'))

    try:
        sections: list[tuple[list[str], list[Path]]] = []
        collect_sections(root, root, sections)
        emit_sections(doc, sections, root, state, tmp_dir, total)
    finally:
        # Чистим временные нормализованные копии
        try:
            shutil.rmtree(tmp_dir, ignore_errors=True)
        except Exception:
            pass

    # Сохранение
    cp(f"\n{'─' * 70}", C.BOLD)
    cp("ШАГ 3. Сохранение docx", C.BOLD)
    cp("─" * 70, C.BOLD)

    try:
        doc.save(str(out_path))
    except PermissionError:
        cp(f"⚠ Не могу записать {out_path.name} — возможно, файл открыт в Word.",
           C.RED)
        cp("Закройте его и запустите скрипт заново.", C.RED)
        return
    except Exception as e:
        cp(f"✗ Ошибка при сохранении: {e}", C.RED)
        return

    # Обновляем поля (TOC + SEQ) через Word, чтобы файл сразу читался без F9
    cp("Обновляю оглавление и номера рисунков через MS Word…", C.CYAN)
    fields_updated = update_fields_via_word(out_path)
    if fields_updated:
        cp("  ✓ Поля обновлены.", C.GREEN)

    cp(f"\n{'═' * 70}", C.BOLD)
    cp("ГОТОВО", C.BOLD + C.GREEN)
    cp(f"{'═' * 70}", C.BOLD)
    cp(f"  Файл:                {out_path}", C.CYAN)
    cp(f"  Размещено фото:      {state.total_processed}", C.GREEN)


if __name__ == '__main__':
    try:
        main()
    except KeyboardInterrupt:
        cp("\n\nПрервано пользователем (Ctrl+C).", C.RED)
        sys.exit(1)
    except Exception as e:
        cp(f"\nНепредвиденная ошибка: {e}", C.RED)
        import traceback
        traceback.print_exc()
        sys.exit(1)
