#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""10_make_tz1_for_contract_ekcelo_v1 — формирователь Технического задания №1.

Скрипт собирает из папки проекта (с произвольной вложенностью):
  • выписки ЕГРН   (XML, PDF) — об объектах недвижимости;
  • выписки ЕГРЮЛ  (PDF, XML) — о юридических лицах;
  • выписки ЕГРИП  (PDF, XML) — об ИП;
  • ОСВ            (XLSX)     — оборотно-сальдовые ведомости 1С (счёт 01);
  • реестры        (XLSX/PDF) — списки недвижимости / основных средств;

— и формирует DOCX «TZ1<YYYYMMDD_HHMMSS>.docx» с единственной таблицей:

  «Таблица 1. Список объектов для исследования»
      объекты, у которых найдена связь актив↔пассив
      (т.е. для актива из ОСВ/реестра найден правообладатель в ЕГРН,
       или объект ЕГРН однозначно сопоставлен бухгалтерской позиции).

Колонки таблицы:
  №  |  Кадастровый № / Инв.№  |  Наименование актива  |  Адрес  |  Бенефициар

Колонка «Наименование актива» синтезируется из четырёх характеристик:
  <Тип>; <Параметр>; <Назначение>; <Наименование>
где
  Тип        — земельный участок / здание / помещение / сооружение /
               машино-место / ОНС / оборудование;
  Параметр   — площадь (кв.м), протяжённость (м), глубина (м), объём (куб.м);
  Назначение — для участка: категория земель + виды разрешённого использования;
               для здания/сооружения/помещения/ОНС: назначение;
  Наименование — «Наименование» из выписки/ОСВ (как самостоятельный реквизит).

«Бенефициар» — непосредственный правообладатель из ЕГРН (наименование + ИНН).

Запуск (VSCode terminal, Win10, русская локаль):
    python parser/scripts/pirushin_sosn_rocha_10_make_tz1_for_contract_ekcelo_v1.py

Загрузка данных — гибридная:
  1. Если в <project>/_data/structure.json есть актуальный снапшот (сборка 052) —
     берём оттуда (cadastre_objects + beneficiaries).
  2. Дополнительно (или при отсутствии structure.json) — рекурсивный
     обход и парсинг сырых файлов через egrn_parser/parsers/* (если модули
     доступны) либо через лёгкие fallback-ридеры.
"""
from __future__ import annotations

import json
import re
import sys
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Iterable

SCRIPTS = Path(__file__).resolve().parent
sys.path.insert(0, str(SCRIPTS.parent))          # parser/
sys.path.insert(0, str(SCRIPTS.parent.parent))   # ekcelo/


# ─── Console UI (Win10 VSCode terminal, русская локаль) ─────────────────────

class C:
    G = "\033[92m"; R = "\033[91m"; Y = "\033[93m"; CY = "\033[96m"
    B = "\033[1m"; X = "\033[0m"


def cp(t: str = "", c: str = "") -> None:
    print(f"{c}{t}{C.X}" if c else t)


def ask(prompt: str, default: str = "") -> str:
    suffix = f" [{default}]" if default else ""
    return input(f"{prompt}{suffix}: ").strip() or default


def ask_path(prompt: str, must_exist: bool = True,
              default: str | None = None) -> Path:
    while True:
        raw = ask(prompt, default or "")
        if not raw:
            cp("  Путь не может быть пустым.", C.Y)
            continue
        # Принимаем как Windows-стиль («E:\\Code\\...») так и POSIX
        p = Path(raw.strip(' "\'')).expanduser()
        try:
            p = p.resolve()
        except OSError as e:
            cp(f"  Не получилось разобрать путь: {e}", C.Y)
            continue
        if must_exist and not p.exists():
            cp(f"  Путь не существует: {p}", C.Y)
            continue
        return p


def ask_yn(prompt: str, default: bool = True) -> bool:
    suffix = " (Y/n): " if default else " (y/N): "
    while True:
        ans = input(prompt + suffix).strip().lower()
        if not ans:
            return default
        if ans in ("y", "yes", "д", "да"):
            return True
        if ans in ("n", "no", "н", "нет"):
            return False
        cp("  Введите 'y' / 'n' (или 'д' / 'н').", C.Y)


# ─── Модель данных ──────────────────────────────────────────────────────────

@dataclass
class Asset:
    """Сводная запись об объекте (актив).

    Идентификация: cad_number ИЛИ inv_number (бух. учёт).

    Колонка «Наименование актива» в ТЗ-1 синтезируется из 4 полей через
    property `name`:  <kind>; <param>; <purpose>; <own_name>.
    """
    cad_number: str | None = None       # кадастровый номер
    inv_number: str | None = None       # инвентарный № (бух. учёт)
    kind: str = ""                      # тип: земельный участок / здание / …
    param: str = ""                     # площадь / протяжённость / объём
    purpose: str = ""                   # назначение (или категория + ВРИ)
    own_name: str = ""                  # «Наименование» как реквизит
    address: str = ""                   # адрес местонахождения
    holders: list[tuple[str | None, str | None]] = field(default_factory=list)
    source: str = ""                    # «ЕГРН (XML)», «ОСВ», «Реестр», …
    source_path: str = ""               # путь к файлу-источнику

    def add_holder(self, name: str | None, inn: str | None) -> None:
        """Добавить правообладателя с дедупликацией (по ИНН → по имени)."""
        name = (name or "").strip() or None
        inn = (inn or "").strip() or None
        if not name and not inn:
            return
        for ex_name, ex_inn in self.holders:
            if inn and ex_inn and inn == ex_inn:
                return
            if not inn and not ex_inn and name and ex_name \
                    and name.lower() == ex_name.lower():
                return
        self.holders.append((name, inn))

    @property
    def name(self) -> str:
        """Синтезированное наименование актива для колонки таблицы."""
        parts = [p for p in (self.kind, self.param, self.purpose,
                              self.own_name) if p]
        # Не дублируем own_name, если он уже включён в purpose
        if (self.purpose and self.own_name
                and self.own_name.lower() in self.purpose.lower()):
            parts = [p for p in (self.kind, self.param, self.purpose) if p]
        return "; ".join(parts) if parts else "—"

    def key(self) -> str:
        if self.cad_number:
            return f"cn::{self.cad_number}"
        if self.inv_number:
            return f"inv::{self.inv_number}"
        return (f"own::{self.own_name.strip().lower()}::"
                f"{self.address.strip().lower()}")

    def has_link(self) -> bool:
        """Есть ли связь актив↔пассив (известен хотя бы один правообладатель)."""
        return bool(self.holders)

    def merge_from(self, other: "Asset") -> None:
        """Слить недостающие поля из `other` (без перетирания)."""
        for fld in ("cad_number", "inv_number", "kind", "param", "purpose",
                    "own_name", "address"):
            if not getattr(self, fld) and getattr(other, fld):
                setattr(self, fld, getattr(other, fld))
        for hn, hi in other.holders:
            self.add_holder(hn, hi)
        # Источник: накапливаем все
        if other.source and other.source not in self.source:
            self.source = (self.source + "; " + other.source).strip("; ")
        if other.source_path and other.source_path != self.source_path:
            sp = self.source_path
            self.source_path = (sp + "; " + other.source_path).strip("; ") if sp \
                else other.source_path


@dataclass
class Beneficiary:
    """Сведения о юр.лице / ИП из ЕГРЮЛ / ЕГРИП."""
    inn: str | None = None
    ogrn: str | None = None
    name: str = ""
    source: str = ""
    source_path: str = ""


@dataclass
class Accumulator:
    assets: dict[str, Asset] = field(default_factory=dict)
    beneficiaries: dict[str, Beneficiary] = field(default_factory=dict)

    def add_asset(self, a: Asset) -> None:
        k = a.key()
        if k in self.assets:
            self.assets[k].merge_from(a)
        else:
            self.assets[k] = a

    def add_beneficiary(self, b: Beneficiary) -> None:
        k = b.inn or b.ogrn or b.name.strip().lower()
        if not k:
            return
        if k in self.beneficiaries:
            ex = self.beneficiaries[k]
            for fld in ("inn", "ogrn", "name"):
                if not getattr(ex, fld) and getattr(b, fld):
                    setattr(ex, fld, getattr(b, fld))
        else:
            self.beneficiaries[k] = b


# ─── Фаза 1. Гибридная загрузка из _data/structure.json ─────────────────────

def load_from_structure(root: Path, acc: Accumulator) -> bool:
    """Если в <project>/_data/ есть structure[_*].json — заполняем acc.

    Возвращает True, если структура найдена и прочитана.
    """
    candidates = []
    data_dir = root / "_data"
    if data_dir.exists():
        candidates.append(data_dir / "structure.json")
        candidates.extend(sorted(data_dir.glob("structure_*.json")))
        candidates.extend(sorted(data_dir.glob("enriched*.json")))
    structure = None
    for p in candidates:
        if p.exists():
            try:
                structure = json.loads(p.read_text(encoding="utf-8"))
                cp(f"  ✓ Прочитан snapshot: {p.relative_to(root)}", C.G)
                break
            except Exception as e:  # noqa: BLE001
                cp(f"  ⚠ Не удалось прочитать {p.name}: {e}", C.Y)
    if not structure:
        return False

    # Бенефициары
    for bk, b in (structure.get("beneficiaries") or {}).items():
        attrs = b.get("attrs") if isinstance(b, dict) and isinstance(
            b.get("attrs"), dict) else (b if isinstance(b, dict) else {})
        ben = Beneficiary(
            inn=str(attrs.get("ИНН") or b.get("ИНН") or "").strip() or None,
            ogrn=str(attrs.get("ОГРН") or b.get("ОГРН") or "").strip() or None,
            name=(attrs.get("Полное наименование")
                  or attrs.get("Краткое наименование")
                  or b.get("Полное наименование") or bk).strip(),
            source="structure.json", source_path=str(data_dir),
        )
        acc.add_beneficiary(ben)

    # Объекты
    cads = structure.get("cadastre_objects") or []
    for c in cads:
        kind = _normalize_kind(c.get("object_type") or "")
        # параметр: площадь либо протяжённость
        area_v = c.get("area")
        param = ""
        if area_v not in (None, "", 0):
            param = f"{_fmt_num(str(area_v))} кв.м"
        elif c.get("length"):
            param = f"протяжённость {_fmt_num(str(c['length']))} м"
        # назначение: для участка — категория + ВРИ, иначе — purpose/назначение
        if kind == "земельный участок":
            cat = (c.get("land_category") or c.get("category") or "").strip()
            vri = (c.get("permitted_use") or c.get("vri") or "").strip()
            if cat and vri:
                purpose = f"{cat}; ВРИ: {vri}"
            else:
                purpose = cat or (f"ВРИ: {vri}" if vri else "")
        else:
            purpose = (c.get("purpose") or c.get("designation") or "").strip()
        a = Asset(
            cad_number=(c.get("cadastral_number") or "").strip() or None,
            inv_number=(c.get("inv_number") or "").strip() or None,
            kind=kind,
            param=param,
            purpose=purpose,
            own_name=(c.get("name") or c.get("object_name") or "").strip(),
            address=(c.get("address") or c.get("text_descriptor") or "").strip(),
            source="structure.json", source_path=str(data_dir),
        )
        # Правообладатели: ВСЕ из rights/rightholders
        for r in (c.get("rights") or c.get("rightholders") or []):
            if not isinstance(r, dict):
                continue
            hn = (r.get("beneficiary_name") or r.get("holder_name") or "")
            hi = str(r.get("beneficiary_inn") or r.get("inn") or "")
            a.add_holder(hn, hi)
        acc.add_asset(a)
    return bool(cads or structure.get("beneficiaries"))


# ─── Фаза 2. Сырые ридеры (fallback) ────────────────────────────────────────

CAD_RE = re.compile(r"\b\d{2}:\d{2}:\d{6,7}:\d{1,7}\b")
INN_RE = re.compile(r"\b(?:\d{10}|\d{12})\b")
OGRN_RE = re.compile(r"\b(?:\d{13}|\d{15})\b")
EGRN_MARKERS = (
    "выписка из единого государственного реестра недвижимости",
    "сведения о характеристиках объекта недвижимости",
    "сведения об основных характеристиках",
)
EGRUL_MARKERS = (
    "сведения о юридическом лице",
    "единый государственный реестр юридических лиц",
)
EGRIP_MARKERS = (
    "сведения об индивидуальном предпринимателе",
    "единый государственный реестр индивидуальных предпринимателей",
)
OSV_MARKERS = ("оборотно-сальдовая ведомость", "осв",
               "01.01", "01.03", "01.к", "01.К", "счет 01")


def _read_pdf_text(p: Path, max_pages: int | None = None) -> str:
    """Читаем ВЕСЬ PDF как текст (max_pages=None) — иначе многообъектные
    выписки ЕГРН обрезаются и теряются объекты со 2-го по N-ный.

    Параметр `max_pages` оставлен для диагностики/отладки; по умолчанию
    лимита нет. Тихий graceful-fail при отсутствии pdfplumber или ошибке.
    """
    try:
        import warnings
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            import pdfplumber  # type: ignore[import]
    except ImportError:
        return ""
    out: list[str] = []
    try:
        with pdfplumber.open(str(p)) as pdf:
            pages = pdf.pages if max_pages is None else pdf.pages[:max_pages]
            for page in pages:
                try:
                    t = page.extract_text() or ""
                except Exception:  # noqa: BLE001
                    t = ""
                out.append(t)
    except Exception:  # noqa: BLE001
        return ""
    return "\n".join(out)


def _classify_text(text_lower: str) -> str:
    if any(m in text_lower for m in EGRN_MARKERS):
        return "egrn"
    if any(m in text_lower for m in EGRUL_MARKERS):
        return "egrul"
    if any(m in text_lower for m in EGRIP_MARKERS):
        return "egrip"
    return "other"


# Мусорные подстроки, которые НЕЛЬЗЯ принимать за значение поля
# (это служебная разметка PDF-выписки ЕГРН и стандартные «пустые» значения).
_GARBAGE_SUBSTR = (
    "лист №", "лист№", "всего листов", "всего разделов",
    "раздел ", "раздела ",
    "сведения отсутствуют", "данные отсутствуют",
    "не установлен", "не установлено", "не установлена",
    "отсутствует", "отсутствуют",
    "выписка", "росреестр", "роскадастр",
)


def _clean(s: str) -> str:
    """Срез пробелов/знаков пунктуации, отбраковка служебных строк."""
    if not s:
        return ""
    s = s.strip().strip(",;:.|").strip()
    if not s:
        return ""
    low = s.lower()
    for g in _GARBAGE_SUBSTR:
        if g in low:
            return ""
    return s


# Распознавание типа объекта недвижимости
_KIND_MAP: tuple[tuple[str, str], ...] = (
    ("земельный участок",                          "земельный участок"),
    ("здание",                                     "здание"),
    ("помещение",                                  "помещение"),
    ("сооружение",                                 "сооружение"),
    ("машино-место",                               "машино-место"),
    ("единый недвижимый комплекс",                 "единый недвижимый комплекс"),
    ("объект незавершенного строительства",        "ОНС"),
    ("объект незавершённого строительства",        "ОНС"),
    ("объект незавершенный строительством",        "ОНС"),
)


_KIND_ALIASES: dict[str, str] = {
    "land": "земельный участок", "zu": "земельный участок",
    "building": "здание", "build": "здание",
    "room": "помещение", "flat": "помещение",
    "structure": "сооружение", "construction": "сооружение",
    "parking": "машино-место",
    "ons": "ОНС", "unfinished": "ОНС",
    "equipment": "оборудование",
    "principal_unregistered": "",
}


def _normalize_kind(raw: str) -> str:
    """Приведение `object_type` (из structure.json) к человекочитаемому типу."""
    if not raw:
        return ""
    low = raw.strip().lower()
    if low in _KIND_ALIASES:
        return _KIND_ALIASES[low]
    # Возможно уже по-русски — проверяем через _detect_kind
    return _detect_kind(raw) or raw


def _detect_kind(text: str) -> str:
    """Тип объекта: ищем первое вхождение из словаря."""
    low = text.lower()
    best_pos = len(low) + 1
    best_kind = ""
    for needle, label in _KIND_MAP:
        i = low.find(needle)
        if 0 <= i < best_pos:
            best_pos = i
            best_kind = label
    return best_kind


# Числовой токен с возможной запятой и группировкой
_NUM = r"\d+(?:[ \xa0]\d{3})*(?:[.,]\d+)?"


# ─── Универсальный мультистрочный экстрактор полей ─────────────────────────
#
# Значения полей в PDF-выписке ЕГРН часто переносятся на несколько строк
# («Наименование» с длинным перечнем комнат, «Особые отметки», «Местоположение»
# и т.п.). Простое `[^\n]+` обрезает значение на первом переносе.
#
# Стратегия: знаем перечень ВСЕХ возможных меток полей; значение текущего
# поля — всё от его метки до следующей известной метки в начале строки или конца
# текста. Внутренние переносы строк схлопываются в пробелы.

_FIELD_LABEL_PATTERNS: tuple[str, ...] = (
    r"Номер кадастрового квартала",
    r"Дата присвоения кадастрового номера",
    r"Ранее присвоенный государственный учетный номер",
    r"Местоположение",
    r"Адрес",
    r"Площадь",
    r"Назначение",
    r"Наименование",
    r"Номер, тип этажа",
    r"Этаж",
    r"Вид жилого помещения",
    r"Кадастровая стоимость",
    r"Кадастровые номера иных объектов недвижимости",
    r"Виды разреш[ёе]нного использования",
    r"Категория земель",
    r"Сведения об отнесении",
    r"Статус записи об объекте",
    r"Особые отметки",
    r"Правообладатель",
    r"Сведения о возможности",
    r"Вид, номер, дата и время государственной регистрации",
    r"Документы-основания",
    r"Сведения об осуществлении государственной регистрации",
    r"Ограничение прав",
    r"Заявленные в судебном порядке",
    r"Сведения о возражении",
    r"Сведения о наличии решения",
    r"Сведения о невозможности",
    r"Правопритязания",
    r"Протяж[ёе]нность",
    r"Глубина",
    r"Высота",
    r"Объ[её]м",
    r"Степень готовности",
    r"Кадастровый номер",
    r"Лист\s*№",
    r"Раздел\s*\d",
    r"ДОКУМЕНТ ПОДПИСАН",
    r"Сертификат",
    r"полное наименование",
    r"Получатель выписки",
    r"Сведения, необходимые",
    r"Филиал публично-правовой",
    r"Выписка из Единого",
    r"На основании запроса",
    r"М\.П\.",
    r"\d+(?:\.\d+)?\s+Правообладатель",
    r"\d+(?:\.\d+)?\s+Вид, номер",
    r"\d+(?:\.\d+)?\s+Документы",
    r"\d+(?:\.\d+)?\s+Ограничение",
    r"\d+(?:\.\d+)?\s+Сведения",
    r"\d+(?:\.\d+)?\s+Заявленные",
    r"\d+(?:\.\d+)?\s+Правопритязания",
)

_NEXT_LABEL_LA = (
    r"(?=\n\s*(?:" + "|".join(_FIELD_LABEL_PATTERNS) + r")\b|\Z)"
)


def _read_field(text: str, label_pat: str) -> str:
    """Извлекает мультистрочное значение поля.

    `label_pat` — regex-фрагмент с меткой (без двоеточия и без якоря).
    Значение — от метки до следующей известной метки или конца текста;
    внутренние переносы склеиваются пробелами; результат прогоняется
    через `_clean()` (фильтр служебной разметки).
    """
    # После метки часто идёт уточнение/единица измерения и двоеточие, например:
    # «Площадь, м2: 727.7», «Кадастровая стоимость, руб: …». Поглощаем любые
    # non-newline-non-colon символы до обязательного «:» (или допускаем простой
    # пробельный разделитель, если двоеточия нет).
    pat = (
        rf"(?:^|\n)\s*{label_pat}"
        rf"(?:[^\n:]*?:|\s)\s*"
        rf"(.+?){_NEXT_LABEL_LA}"
    )
    m = re.search(pat, text, flags=re.DOTALL | re.IGNORECASE)
    if not m:
        return ""
    val = m.group(1)
    val = re.sub(r"\s*\n\s*", " ", val)
    val = re.sub(r"\s{2,}", " ", val).strip().strip(",;:")
    return _clean(val)


def _read_number(text: str, label_pat: str) -> str:
    """Берёт значение поля и возвращает первое число из него (форматированное)."""
    val = _read_field(text, label_pat)
    if not val:
        return ""
    m = re.search(rf"({_NUM})", val)
    return _fmt_num(m.group(1)) if m else ""


def _fmt_num(raw: str) -> str:
    """«1 234,56» → «1234,56» (убираем NBSP/пробелы внутри групп)."""
    return re.sub(r"[ \xa0]+", "", raw).strip()


def _extract_address(text: str) -> str:
    return _read_field(text, r"(?:Адрес|Местоположение)")


def _extract_param(text: str, kind: str) -> str:
    """Главный размерный параметр в зависимости от типа объекта."""
    found: list[str] = []
    if n := _read_number(text, r"Площадь"):
        found.append(f"{n} кв.м")
    if n := _read_number(text, r"Протяж[ёе]нность"):
        found.append(f"протяжённость {n} м")
    if n := _read_number(text, r"Глубина"):
        found.append(f"глубина {n} м")
    if n := _read_number(text, r"Высота"):
        found.append(f"высота {n} м")
    if n := _read_number(text, r"Объ[её]м"):
        found.append(f"объём {n} куб.м")
    if kind == "ОНС":
        if n := _read_number(text, r"Степень\s+готовности"):
            found.append(f"готовность {n}%")
    return ", ".join(found)


def _extract_purpose(text: str, kind: str) -> str:
    """Назначение — по-разному для участка и для прочих типов."""
    if kind == "земельный участок":
        cat = _read_field(text, r"Категория\s+земель")
        vri = _read_field(text, r"Вид(?:ы)?\s+разреш[ёе]нного\s+использования")
        if cat and vri:
            return f"{cat}; ВРИ: {vri}"
        return cat or (f"ВРИ: {vri}" if vri else "")
    return _read_field(
        text,
        r"Назначение(?!\s+земель|\s+участка|\s+жилого)",
    )


def _extract_own_name(text: str) -> str:
    """«Наименование» как самостоятельный реквизит выписки."""
    return _read_field(
        text,
        r"Наименование(?!\s+правооблад|\s+органа|\s+документа|\s+по\s)",
    )


# Метка «Кадастровый номер: NN:NN:NNNNNNN:NN» — стандартный заголовок
# объекта в выписке ЕГРН. Может появляться многократно (в шапке/подвале
# каждой страницы); многообъектная выписка содержит ПОСЛЕДОВАТЕЛЬНОСТЬ
# разных КН — по их «runs» и режется текст на объектные блоки.
_CAD_LABEL_RE = re.compile(
    r"Кадастровый\s+номер[\s:]+(\d{2}:\d{2}:\d{6,7}:\d+)",
    re.IGNORECASE,
)


def _distinct_cads(text: str) -> list[str]:
    """Все уникальные КН из меток «Кадастровый номер: …» в порядке появления.

    В выписке метка повторяется на каждой странице (в шапке/подвале) — берём
    только различные. В типичной выписке (один объект) возвращается ровно
    один КН.
    """
    seen: set[str] = set()
    out: list[str] = []
    for m in _CAD_LABEL_RE.finditer(text):
        cad = m.group(1)
        if cad not in seen:
            seen.add(cad)
            out.append(cad)
    return out


# Блок «1.N <Наименование> ИНН <код>» — стандартная разбивка правообладателей
# в разделе 2 выписки ЕГРН (см. egrn_parser/parsers/pdf_parser.py:704).
_HOLDER_BLOCK_RE = re.compile(
    r"1\.\d+\s+([^\n]+?)\s+ИНН[\s:]*?(\d{10}|\d{12})",
    re.IGNORECASE,
)
_HOLDER_FALLBACK_RE = re.compile(
    r"Правообладател[ьи][\s\S\(\)]*?:[\s\S]*?([А-ЯA-Z][^\n]+?)\s+ИНН[\s:]*?(\d{10}|\d{12})",
    re.IGNORECASE,
)


def _extract_holders(text: str) -> list[tuple[str, str]]:
    """Все правообладатели из раздела 2 (или fallback одиночный)."""
    res: list[tuple[str, str]] = []
    seen_inns: set[str] = set()
    for m in _HOLDER_BLOCK_RE.finditer(text):
        name, inn = m.group(1).strip().strip(",;"), m.group(2)
        if inn in seen_inns:
            continue
        seen_inns.add(inn)
        res.append((name, inn))
    if not res:
        m = _HOLDER_FALLBACK_RE.search(text)
        if m:
            res.append((m.group(1).strip().strip(",;"), m.group(2)))
    return res


def _ingest_egrn_block(p: Path, cad: str, block: str, acc: Accumulator,
                        fmt: str, kind_hint: str = "") -> None:
    """Один объект ЕГРН → Asset + бенефициары. Используется и для PDF, и для XML.

    `fmt` — «PDF» / «XML» (для атрибута `source`).
    `kind_hint` — заранее известный тип (например, из корневого тега XML);
    если пуст, тип угадывается из текста.
    """
    kind = kind_hint or _detect_kind(block)
    holders = _extract_holders(block)
    a = Asset(
        cad_number=cad,
        kind=kind,
        param=_extract_param(block, kind),
        purpose=_extract_purpose(block, kind),
        own_name=_extract_own_name(block),
        address=_extract_address(block),
        source=f"ЕГРН ({fmt})", source_path=str(p),
    )
    for hn, hi in holders:
        a.add_holder(hn, hi)
    acc.add_asset(a)
    for hn, hi in holders:
        acc.add_beneficiary(Beneficiary(
            inn=hi, name=hn,
            source=f"ЕГРН ({fmt})", source_path=str(p),
        ))


def scan_pdf_egrn(p: Path, acc: Accumulator) -> None:
    text = _read_pdf_text(p)
    if not text:
        return
    cls = _classify_text(text.lower())
    if cls == "egrn":
        # Берём ВЕСЬ текст выписки и работаем с ним за один проход:
        # поля (Наименование, Площадь, Назначение, …) могут располагаться
        # ДО первой метки «Кадастровый номер:» (например, в шапке-таблице
        # PDF), поэтому делить текст по меткам КН — ошибка.
        cads = _distinct_cads(text)
        if not cads:
            cad_fallback = next(iter(CAD_RE.findall(text)), None)
            if not cad_fallback:
                return
            cads = [cad_fallback]
        primary_cad = cads[0]
        _ingest_egrn_block(p, primary_cad, text, acc, fmt="PDF")
        if len(cads) > 1:
            cp(f"  ⚠ {p.name}: многообъектная выписка ЕГРН, КН={cads}; "
               f"распарсен только {primary_cad}, остальные требуют отдельной "
               f"обработки.", C.Y)
    elif cls in ("egrul", "egrip"):
        # Заголовок юр.лица — обычно «Полное наименование: …»
        m = re.search(r"Полное наименование[\s:]+([^\n]+)", text,
                      flags=re.IGNORECASE)
        full_name = m.group(1).strip() if m else ""
        inn_m = INN_RE.search(text); inn = inn_m.group(0) if inn_m else None
        ogrn_m = OGRN_RE.search(text); ogrn = ogrn_m.group(0) if ogrn_m else None
        if full_name or inn or ogrn:
            acc.add_beneficiary(Beneficiary(
                inn=inn, ogrn=ogrn, name=full_name,
                source=f"{cls.upper()} (PDF)", source_path=str(p),
            ))


def scan_xml(p: Path, acc: Accumulator) -> None:
    """Парсим XML; если это ЕГРН-выписка — извлекаем КН/адрес/правообладателя."""
    from xml.etree import ElementTree as ET
    try:
        tree = ET.parse(str(p))
    except Exception:  # noqa: BLE001
        return
    root = tree.getroot()
    tag = root.tag.split("}")[-1].lower()
    if "extract_about_property" not in tag and "egrul" not in tag and "egrip" not in tag:
        return

    def _txt(el) -> str:
        return "".join(el.itertext()) if el is not None else ""

    full_text = _txt(root)
    if "extract_about_property" in tag:
        labelled = "\n".join(
            e.tag.split("}")[-1] + ": " + (e.text or "")
            for e in root.iter() if e.text
        )
        # Тип объекта — приоритетно из имени корневого тега XML.
        kind_from_tag = {
            "extract_about_property_land":         "земельный участок",
            "extract_about_property_build":        "здание",
            "extract_about_property_room":         "помещение",
            "extract_about_property_construction": "сооружение",
            "extract_about_property_parking":      "машино-место",
            "extract_about_property_ons":          "ОНС",
        }.get(tag, "")
        # XML-выписка ЕГРН — один объект (по корневому тегу).
        merged = labelled or full_text
        cads = _distinct_cads(merged)
        cad = cads[0] if cads else next(iter(CAD_RE.findall(merged)), None)
        if not cad:
            return
        _ingest_egrn_block(p, cad, merged, acc, fmt="XML",
                            kind_hint=kind_from_tag)
    else:
        m = re.search(r"Полное наименование[\s:]+([^\n]+)", full_text,
                      flags=re.IGNORECASE)
        full_name = m.group(1).strip() if m else ""
        inn_m = INN_RE.search(full_text); inn = inn_m.group(0) if inn_m else None
        ogrn_m = OGRN_RE.search(full_text); ogrn = ogrn_m.group(0) if ogrn_m else None
        if full_name or inn or ogrn:
            acc.add_beneficiary(Beneficiary(
                inn=inn, ogrn=ogrn, name=full_name,
                source=("ЕГРЮЛ (XML)" if "egrul" in tag else "ЕГРИП (XML)"),
                source_path=str(p),
            ))


def scan_xlsx(p: Path, acc: Accumulator) -> None:
    """ОСВ (счёт 01) или реестр недвижимости. Лёгкий ридер по заголовкам."""
    try:
        import openpyxl  # type: ignore[import]
    except ImportError:
        return
    try:
        wb = openpyxl.load_workbook(str(p), data_only=True, read_only=True)
    except Exception:  # noqa: BLE001
        return

    for ws in wb.worksheets:
        # Сначала собираем «шапку» (первые 12 строк)
        head_rows: list[list[str]] = []
        for ri, row in enumerate(ws.iter_rows(values_only=True)):
            head_rows.append([("" if c is None else str(c)) for c in row])
            if ri >= 11:
                break
        head_text = "\n".join(" | ".join(r) for r in head_rows).lower()
        is_osv = any(m.lower() in head_text for m in OSV_MARKERS)
        is_registry = ("кадастр" in head_text or "инв" in head_text
                       or "наимен" in head_text) and ("адрес" in head_text)
        if not (is_osv or is_registry):
            continue

        # Тип объекта по имени листа (для шаблона Assets)
        sheet_kind = _sheet_kind(ws.title or "")

        header_row_idx, header_cells = _find_header_row(head_rows)
        if header_row_idx is None:
            continue
        col_map = _map_columns(header_cells)
        if not col_map:
            continue

        for row in ws.iter_rows(min_row=header_row_idx + 2, values_only=True):
            if not any(row):
                continue
            cells = [("" if c is None else str(c)).strip() for c in row]
            def _g(name: str) -> str:
                idx = col_map.get(name)
                return cells[idx] if idx is not None and idx < len(cells) else ""
            cad = _g("cad")
            inv = _g("inv")
            own_name = _g("name")
            addr = _g("address")
            owner = _g("owner")
            area = _g("area")
            category = _g("category")
            vri = _g("vri")
            purpose_txt = _g("purpose")
            if not (cad or inv or own_name):
                continue
            # Если КН встроен в текст «наименования» — выкусываем
            if not cad:
                m = CAD_RE.search(" ".join(cells))
                if m:
                    cad = m.group(0)
            owner_inn = None
            inn_m = INN_RE.search(owner) if owner else None
            if inn_m:
                owner_inn = inn_m.group(0)

            kind = sheet_kind
            if not kind and cad:
                # Эвристика по 2-му сегменту КН (тип объекта): но достоверно
                # не определить — оставляем пустым, синтез всё равно сложится.
                kind = ""

            param = f"{_fmt_num(area)} кв.м" if area else ""
            purpose = ""
            if kind == "земельный участок":
                if category and vri:
                    purpose = f"{_clean(category)}; ВРИ: {_clean(vri)}"
                elif category:
                    purpose = _clean(category)
                elif vri:
                    purpose = f"ВРИ: {_clean(vri)}"
            else:
                purpose = _clean(purpose_txt)

            a = Asset(
                cad_number=cad or None,
                inv_number=inv or None,
                kind=kind,
                param=param,
                purpose=purpose,
                own_name=_clean(own_name),
                address=_clean(addr),
                source=("ОСВ (XLSX)" if is_osv else "Реестр (XLSX)"),
                source_path=str(p),
            )
            if owner or owner_inn:
                a.add_holder(owner, owner_inn)
            acc.add_asset(a)


def _sheet_kind(sheet_name: str) -> str:
    """Тип объекта, выводимый из имени листа Excel-реестра."""
    low = sheet_name.lower()
    if "земельн" in low or "зу" == low.strip():
        return "земельный участок"
    if "помещ" in low:
        return "помещение"
    if "машино" in low:
        return "машино-место"
    if ("здани" in low and "сооруж" not in low):
        return "здание"
    if "сооруж" in low and "здани" not in low:
        return "сооружение"
    if "оборуд" in low:
        return "оборудование"
    return ""


def _find_header_row(rows: list[list[str]]) -> tuple[int | None, list[str]]:
    """Ищем строку, где встречаются «наимен» И («кадастр» ИЛИ «инв»)."""
    for i, r in enumerate(rows):
        joined = " | ".join(r).lower()
        if "наимен" in joined and ("кадастр" in joined or "инв" in joined
                                   or "адрес" in joined):
            return i, r
    return None, []


def _map_columns(header_cells: list[str]) -> dict[str, int]:
    """Маппинг {name|address|cad|inv|owner|area|category|vri|purpose}."""
    m: dict[str, int] = {}
    for idx, raw in enumerate(header_cells):
        h = raw.lower()
        if "name" not in m and "наимен" in h and "категори" not in h:
            m["name"] = idx
        if "address" not in m and "адрес" in h:
            m["address"] = idx
        if "cad" not in m and "кадастр" in h:
            m["cad"] = idx
        if "inv" not in m and ("инв" in h or "номер по бух" in h):
            m["inv"] = idx
        if "owner" not in m and ("собствен" in h or "правооблад" in h
                                  or "владел" in h or "бенефициар" in h):
            m["owner"] = idx
        if "area" not in m and ("площад" in h or "кв.м" in h or "кв. м" in h):
            m["area"] = idx
        if "category" not in m and "категори" in h and "земел" in h:
            m["category"] = idx
        if "vri" not in m and ("разреш" in h and "использ" in h):
            m["vri"] = idx
        if "purpose" not in m and "назначен" in h:
            m["purpose"] = idx
    return m


# ─── Обход проекта ──────────────────────────────────────────────────────────

SKIP_DIRS = {".git", ".obsidian", "__pycache__", ".venv", "node_modules",
             "_data"}


def walk_project(root: Path) -> Iterable[Path]:
    for p in root.rglob("*"):
        if not p.is_file():
            continue
        if any(part in SKIP_DIRS for part in p.parts):
            continue
        ext = p.suffix.lower()
        if ext in (".pdf", ".xml", ".xlsx"):
            yield p


def scan_raw(root: Path, acc: Accumulator) -> tuple[int, int, int]:
    """Сырое сканирование. Возвращает (pdf, xml, xlsx) количества."""
    n_pdf = n_xml = n_xlsx = 0
    for p in walk_project(root):
        try:
            if p.suffix.lower() == ".pdf":
                scan_pdf_egrn(p, acc); n_pdf += 1
            elif p.suffix.lower() == ".xml":
                scan_xml(p, acc); n_xml += 1
            elif p.suffix.lower() == ".xlsx":
                scan_xlsx(p, acc); n_xlsx += 1
        except Exception as e:  # noqa: BLE001
            cp(f"  ⚠ {p.name}: {e}", C.Y)
    return n_pdf, n_xml, n_xlsx


# ─── Связывание актив↔пассив ────────────────────────────────────────────────

def link_holders(acc: Accumulator) -> None:
    """Дообогащаем правообладателей в `a.holders` по справочнику бенефициаров.

    Если у правообладателя известен ИНН, но имя пустое — берём имя из
    `acc.beneficiaries` (и наоборот).
    """
    if not acc.beneficiaries:
        return
    by_inn = {b.inn: b for b in acc.beneficiaries.values() if b.inn}
    by_name = {b.name.strip().lower(): b for b in acc.beneficiaries.values()
               if b.name}
    for a in acc.assets.values():
        enriched: list[tuple[str | None, str | None]] = []
        for hn, hi in a.holders:
            if hn and not hi:
                b = by_name.get(hn.strip().lower())
                if b and b.inn:
                    hi = b.inn
            if hi and not hn:
                b = by_inn.get(hi)
                if b and b.name:
                    hn = b.name
            enriched.append((hn, hi))
        a.holders = enriched


# Организационно-правовые формы → стандартные сокращения.
# Длинные шаблоны должны сматчиться раньше коротких (см. сортировку ниже).
_LEGAL_FORM_MAP: dict[str, str] = {
    "общество с ограниченной ответственностью": "ООО",
    "публичное акционерное общество": "ПАО",
    "непубличное акционерное общество": "НАО",
    "закрытое акционерное общество": "ЗАО",
    "открытое акционерное общество": "ОАО",
    "акционерное общество": "АО",
    "индивидуальный предприниматель": "ИП",
    "федеральное государственное унитарное предприятие": "ФГУП",
    "федеральное государственное автономное учреждение": "ФГАУ",
    "федеральное государственное бюджетное учреждение": "ФГБУ",
    "государственное унитарное предприятие": "ГУП",
    "муниципальное унитарное предприятие": "МУП",
    "государственное бюджетное учреждение": "ГБУ",
    "муниципальное бюджетное учреждение": "МБУ",
    "автономная некоммерческая организация": "АНО",
    "товарищество собственников жилья": "ТСЖ",
    "товарищество собственников недвижимости": "ТСН",
    "жилищно-строительный кооператив": "ЖСК",
    "сельскохозяйственный производственный кооператив": "СПК",
    "производственный кооператив": "ПК",
    "крестьянское (фермерское) хозяйство": "КФХ",
}
_LEGAL_FORM_SORTED = sorted(_LEGAL_FORM_MAP.items(), key=lambda kv: -len(kv[0]))

_HOLDER_PREFIX_RE = re.compile(
    r"^\s*\(?\s*правообладател[яеьи]+\s*\)?\s*[:\-—]?\s*",
    re.IGNORECASE,
)
_HOLDER_NUM_RE = re.compile(r"^\s*\d+(?:\.\d+)*\s*[)\.\-]?\s*")
_INN_IN_NAME_RE = re.compile(
    r"[,;]?\s*\(?\s*ИНН[\s:]*?\d{10,12}\s*\)?", re.IGNORECASE,
)


def _strip_holder_prefix(name: str) -> str:
    """Срезаем «(правообладатели):», нумерацию «1.2» и хвостовой «(ИНН …)»."""
    s = name or ""
    # Применяем дважды — на случай «правообладатели: 1.2 …»
    for _ in range(2):
        s = _HOLDER_PREFIX_RE.sub("", s)
        s = _HOLDER_NUM_RE.sub("", s)
    s = _INN_IN_NAME_RE.sub("", s)
    return s.strip().strip(",;:")


def _abbreviate_legal_form(name: str) -> str:
    """«Общество с ограниченной ответственностью "Антарес"» → «ООО "Антарес"»."""
    if not name:
        return name
    s = name.strip()
    low = s.lower()
    for full, short in _LEGAL_FORM_SORTED:
        if low.startswith(full):
            rest = s[len(full):].strip()
            return f"{short} {rest}" if rest else short
    return s


def _format_holder(name: str | None, inn: str | None) -> str:
    """Один правообладатель → «ООО \"Х\", ИНН N» (или только имя/только ИНН)."""
    n = _abbreviate_legal_form(_strip_holder_prefix(name or "")) if name else ""
    if n and inn:
        return f"{n}, ИНН {inn}"
    return n or (f"ИНН {inn}" if inn else "")


def beneficiary_cell(a: Asset) -> str:
    """Колонка «Бенефициар»: несколько владельцев — через «; »."""
    parts = [_format_holder(hn, hi) for hn, hi in a.holders]
    parts = [p for p in parts if p]
    return "; ".join(parts) if parts else "—"


# ─── MD-выгрузка ────────────────────────────────────────────────────────────

def _md_escape(s: str) -> str:
    """Экранируем GFM-разделитель ячейки `|` и убираем переносы строк."""
    return (s or "").replace("|", "\\|").replace("\n", " ").replace("\r", " ")


def build_md(table1: list[Asset], out_path: Path, meta: dict) -> Path:
    """Markdown-версия отчёта ТЗ-1 (git-friendly preview).

    Структура зеркалит DOCX: шапка → meta → «Таблица 1» с группировкой
    (по subheading-секциям, по одной таблице на группу) → подписи сторон.
    Зависимостей нет — пишем plain text.
    """
    g = lambda s: _md_escape(s)
    lines: list[str] = []

    # Шапка
    lines.append(
        f"*Приложение 1 (обязательное) к Договору №{meta['doc_num']}-1 "
        f"от {meta['doc_date']}*"
    )
    lines.append("")
    lines.append("# Техническое задание №1")
    lines.append("")
    lines.append(f"Уточнено {meta['time']} {meta['date']}  ")
    lines.append(f"Источник данных: `{meta['root']}`  ")
    lines.append(
        f"Найдено объектов: всего {meta['total']}, "
        f"со связью актив↔пассив: {meta['linked']}  "
    )
    lines.append(
        f"Бенефициаров (ЕГРЮЛ/ЕГРИП/ЕГРН): {meta['beneficiaries']}"
    )
    lines.append("")

    lines.append("## Таблица 1. Список объектов для исследования")
    lines.append("")
    lines.append(
        "В таблицу включены объекты, для которых установлена связь "
        "актив↔пассив: найден непосредственный правообладатель по данным "
        "ЕГРН либо объект однозначно сопоставлен бухгалтерской позиции. "
        "Колонка «Наименование актива» синтезирована по схеме "
        "«Тип; Параметр; Назначение; Наименование»."
    )
    lines.append("")

    if not table1:
        lines.append("_— объектов не найдено —_")
    else:
        grouped = _group_assets_for_table(table1)
        num = 0
        header = ("| № | Кадастровый № / Инв.№ | Наименование актива "
                  "| Адрес местонахождения | Бенефициар |")
        delim  = "|:-:|---|---|---|---|"
        for label, group in grouped:
            lines.append(f"### {label}")
            lines.append("")
            lines.append(header)
            lines.append(delim)
            for a in group:
                num += 1
                ident = g(a.cad_number or a.inv_number or "—")
                name = g(a.name or "—")
                addr = g(a.address or "—")
                benef = g(beneficiary_cell(a))
                lines.append(
                    f"| {num} | {ident} | {name} | {addr} | {benef} |"
                )
            lines.append("")

    # Подписи сторон
    lines.append("## Подписи сторон:")
    lines.append("")
    lines.append("| Исполнитель: | Заказчик: |")
    lines.append("| --- | --- |")
    lines.append("| _________/_________/ | _________/_________/ |")
    lines.append("")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(lines), encoding="utf-8")
    return out_path


# ─── DOCX-выгрузка ──────────────────────────────────────────────────────────

def build_docx(table1: list[Asset], out_path: Path, meta: dict) -> Path:
    from docx import Document  # type: ignore[import]
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for sec in doc.sections:
        sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
        sec.top_margin = sec.bottom_margin = Cm(2.0)
        sec.left_margin = sec.right_margin = Cm(2.0)
    st = doc.styles["Normal"]
    st.font.name = "Arial"; st.font.size = Pt(10)

    # Шапка документа.
    # 1-я строка: «Приложение 1» — курсив, по правому краю;
    # перенос строки (НЕ новый параграф) → «к Договору №…-1 от …г.» — тоже
    # курсив, по правому краю (входит в тот же параграф).
    # Следом: «Техническое задание №1» — обычный шрифт, по центру, БЕЗ
    # горизонтальной полосы (поэтому это не heading — встроенные стили
    # Heading 0/1 рисуют border).
    appendix = doc.add_paragraph()
    appendix.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = appendix.add_run("Приложение 1 (обязательное)")
    r1.italic = True
    r1.add_break()  # line-break внутри параграфа (а не \n)
    r2 = appendix.add_run(
        f"к Договору №{meta['doc_num']}-1 от {meta['doc_date']}"
    )
    r2.italic = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Техническое задание №1")

    meta_p = doc.add_paragraph()
    meta_p.add_run(f"Уточнено {meta['time']} {meta['date']}\n")
    meta_p.add_run(f"Источник данных: {meta['root']}\n")
    meta_p.add_run(
        f"Найдено объектов: всего {meta['total']}, "
        f"со связью актив↔пассив: {meta['linked']}\n"
    )
    meta_p.add_run(f"Бенефициаров (ЕГРЮЛ/ЕГРИП/ЕГРН): {meta['beneficiaries']}")

    _emit_table_block(
        doc,
        heading="Таблица 1. Список объектов для исследования",
        intro=(
            "В таблицу включены объекты, для которых установлена связь "
            "актив↔пассив: найден непосредственный правообладатель по данным "
            "ЕГРН либо объект однозначно сопоставлен бухгалтерской позиции. "
            "Колонка «Наименование актива» синтезирована по схеме "
            "«Тип; Параметр; Назначение; Наименование»."
        ),
        rows=table1,
    )

    _emit_signatures_block(doc)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def _emit_signatures_block(doc) -> None:
    """Подписи сторон: 2-колоночная таблица без рамок.

    Слева  — «Исполнитель:» + строка «_________/_________/»;
    справа — «Заказчик:»    + строка «_________/_________/».
    Заголовок «Подписи сторон:» — слева, без жирного.
    """
    doc.add_paragraph()  # отступ перед блоком
    doc.add_paragraph("Подписи сторон:")
    doc.add_paragraph()  # ещё один отступ

    tbl = doc.add_table(rows=2, cols=2)
    tbl.autofit = False

    tbl.cell(0, 0).text = "Исполнитель:"
    tbl.cell(1, 0).text = "_________/_________/"
    tbl.cell(0, 1).text = "Заказчик:"
    tbl.cell(1, 1).text = "_________/_________/"

    _set_column_widths(tbl, (8.5, 8.5))
    _remove_table_borders(tbl)


def _remove_table_borders(tbl) -> None:
    """Убираем все границы таблицы (для блока подписей)."""
    from docx.oxml.ns import qn  # type: ignore[import]
    from docx.oxml import OxmlElement  # type: ignore[import]
    tbl_el = tbl._tbl
    tbl_pr = tbl_el.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tbl_pr)
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tbl_pr.append(borders)


# Маркеры компонентов адреса (lowercase prefixes).
# «д.» — намеренно НЕ включаем в city_prefixes (коллизия «д. = дом / деревня»);
# для надёжности используем только однозначные city-маркеры.
_CITY_PREFIXES = (
    "г.", "г ", "город ", "пгт", "посёлок", "поселок",
    "п.г.т", "с. ", "село ", "ст-ца", "станица",
)
_STREET_PREFIXES = (
    "ул.", "ул ", "улица",
    "пр.", "пр ", "пр-кт", "проспект", "просп.",
    "пер.", "переулок",
    "бул.", "бульвар", "б-р",
    "пл.", "площадь",
    "ш.", "шоссе",
    "наб.", "набережная",
    "тракт", "линия", "тупик", "проезд", "аллея",
    "мкр.", "мкр ", "микрорайон",
    "тер.", "территория", "квартал",
)
_HOUSE_PREFIXES = (
    "д.", "д ", "дом", "зд.", "здание",
    "стр.", "строение", "влад.", "владение",
)
# Игнорируемые компоненты (квартира/помещение/офис/литер/корпус как
# самостоятельный компонент адреса). «корп. N» внутри одной компоненты
# с домом — сохраняется (часть номера дома); как отдельная компонента —
# отбрасывается.
_DROP_PREFIXES = (
    "кв.", "кв ", "квартира",
    "пом.", "пом ", "помещение",
    "оф.", "офис",
    "комн.", "комната",
    "литер", "лит.",
    "корп.", "корп ", "корпус",  # как отдельная компонента игнорируется
)


def _has_prefix(s: str, prefixes: tuple[str, ...]) -> bool:
    s_low = s.lower().lstrip()
    return any(s_low.startswith(p) for p in prefixes)


def _strip_address_preamble(addr: str) -> str:
    """Срезает вводную фразу Россреестра, если есть.

    Пример: «Местоположение установлено относительно ориентира, расположенного
    в границах участка. Почтовый адрес ориентира: Ростовская обл., г. Ростов-
    на-Дону, ул. Суворова, 52» → «Ростовская обл., г. Ростов-на-Дону,
    ул. Суворова, 52».
    """
    if not addr:
        return ""
    m = re.search(
        r"(?:Почтовый\s+адрес\s+ориентира|Адрес\s+ориентира)[\s:]+(.+)$",
        addr, flags=re.IGNORECASE | re.DOTALL,
    )
    if m:
        return m.group(1).strip()
    return addr


def _looks_like_bare_house(s: str) -> bool:
    """Похоже ли значение на «голый» номер дома без префикса «д./дом» — например,
    «52», «60А», «1/3», «5-7». Исключаем индексы (5+ цифр подряд)."""
    s = s.strip()
    if not s or not s[0].isdigit():
        return False
    if re.match(r"^\d{5,}$", s):
        return False
    # Допускаем цифры, буквы, разделители «/-» и пробел внутри (например
    # «60А корп. 2»).
    return bool(re.match(r"^\d+[\dА-Яа-яA-Za-z/\-\s\.]*$", s))


def _normalize_house(s: str) -> str:
    """«дом №60» / «д.60А» / «владение 7к1» → «д. 60» / «д. 60А» / «д. 7к1»."""
    m = re.match(
        r"\s*(?:дом|здание|строение|владение|зд\.?|стр\.?|влад\.?|д\.?)"
        r"\s*№?\s*(.+)$",
        s, flags=re.IGNORECASE,
    )
    rest = (m.group(1) if m else s).strip().strip(",;")
    return f"д. {rest}" if rest else "д. ?"


def _address_key(addr: str) -> str:
    """Подпись «[город, ]улица, д. NN» для группировки строк таблицы.

    Парсит адрес по компонентам (split по запятой) и выделяет ОТДЕЛЬНО
    город / улицу / дом; флэт/корпус/литер/оф. — игнорируются. Если дом
    не помечен префиксом «д./дом», но в адресе есть «голый» номер после
    улицы — он принимается за дом (например, «ул. Суворова, 52» → «д. 52»).

    Перед парсингом срезается вводная фраза Россреестра «Местоположение
    установлено… Почтовый адрес ориентира: …» (см. `_strip_address_preamble`).

    Если ни улица, ни дом не распознаются — fallback к последним 2-м
    компонентам адреса (но без «дроповых»). Если адрес пуст — спец-метка.
    """
    if not addr:
        return "— без адреса —"
    addr = _strip_address_preamble(addr)
    parts = [p.strip() for p in addr.split(",") if p.strip()]
    if not parts:
        return "— без адреса —"

    city = ""
    street = ""
    house = ""
    for p in parts:
        if not house and _has_prefix(p, _HOUSE_PREFIXES):
            house = _normalize_house(p)
            continue
        if _has_prefix(p, _DROP_PREFIXES):
            continue
        if not city and _has_prefix(p, _CITY_PREFIXES):
            city = p
            continue
        if not street and _has_prefix(p, _STREET_PREFIXES):
            street = p
            continue

    # Bare-house fallback: «..., ул. Суворова, 52» — компонента «52» без
    # префикса распознаётся как дом, ЕСЛИ улица в адресе была найдена.
    if not house and street:
        for p in parts:
            if _looks_like_bare_house(p):
                house = f"д. {p.strip()}"
                break

    pieces = [x for x in (city, street, house) if x]
    if pieces and (street or house):
        return ", ".join(pieces)

    tail = [p for p in parts[-3:] if not _has_prefix(p, _DROP_PREFIXES)]
    return ", ".join(tail[-2:]) if tail else "— без адреса —"


def _kind_order(kind: str) -> int:
    """Порядок типов внутри одной адресной группы.

    1 — земельный участок,
    2 — здание / сооружение / ОНС,
    3 — помещение / машино-место,
    99 — прочее (страховка).
    """
    if kind == "земельный участок":
        return 1
    if kind in ("здание", "сооружение", "ОНС"):
        return 2
    if kind in ("помещение", "машино-место"):
        return 3
    return 99


def _group_assets_for_table(
    assets: list[Asset],
) -> list[tuple[str, list[Asset]]]:
    """Возвращает [(group_label, [Asset, …]), …] в порядке вывода.

    Сначала — адресные группы (отсортированы alphabetic'ом по label'у)
    с внутренним порядком земля → здания/сооружения/ОНС → помещения.
    Объекты без адреса — в группу «— без адреса —» (после адресных).
    В конце — отдельная группа «Оборудование» (БЕЗ группировки по
    адресу, отсортирована по инв.№ → наименованию).
    """
    equipment = [a for a in assets if a.kind == "оборудование"]
    real_estate = [a for a in assets if a.kind != "оборудование"]

    groups: dict[str, list[Asset]] = {}
    for a in real_estate:
        groups.setdefault(_address_key(a.address), []).append(a)

    def _sort_addr(k: str) -> tuple[int, str]:
        return (1 if k == "— без адреса —" else 0, k.lower())

    out: list[tuple[str, list[Asset]]] = []
    for key in sorted(groups.keys(), key=_sort_addr):
        group = sorted(groups[key], key=lambda a: (
            _kind_order(a.kind),
            (a.cad_number or "") + (a.inv_number or ""),
        ))
        out.append((key, group))

    if equipment:
        equipment_sorted = sorted(equipment, key=lambda a: (
            a.inv_number or "", a.own_name or ""
        ))
        out.append(("Оборудование", equipment_sorted))
    return out


def _emit_table_block(doc, heading: str, intro: str, rows: list[Asset]) -> None:
    from docx.enum.table import WD_ALIGN_VERTICAL  # type: ignore[import]
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import]

    doc.add_heading(heading, level=1)
    doc.add_paragraph(intro)
    if not rows:
        doc.add_paragraph("— объектов не найдено —").italic = True
        return

    headers = ["№", "Кадастровый № / Инв.№", "Наименование актива",
               "Адрес местонахождения", "Бенефициар"]
    # Ширины колонок на A4 (зона контента 17 см).
    widths_cm = (1.0, 3.5, 5.5, 3.5, 3.5)  # = 17.0 см
    n_cols = len(headers)
    tbl = doc.add_table(rows=1, cols=n_cols)
    # «Table Grid» — простая сетка без банднинга и цветной шапки. Группы
    # выделяются мягкой серой подложкой только на subheading-row;
    # ячейки данных остаются без фона.
    tbl.style = "Table Grid"
    tbl.autofit = False

    # Шапка таблицы: жирная, центрирование H + V; повторяется на следующих
    # страницах (w:tblHeader).
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        cell.text = h
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in para.runs:
            r.bold = True
    _mark_row_as_header(tbl.rows[0])

    grouped = _group_assets_for_table(rows)
    num = 0
    for label, group in grouped:
        # Заголовок группы: одна строка, объединённая на все колонки,
        # жирный текст по левому краю. Subheading-row короткая —
        # всегда «не разрывать» (cantSplit).
        sub_row = tbl.add_row()
        sub_cell = sub_row.cells[0]
        for i in range(1, n_cols):
            sub_cell.merge(sub_row.cells[i])
        sub_cell.text = label
        sub_para = sub_cell.paragraphs[0]
        sub_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in sub_para.runs:
            r.bold = True
        # Мягкая серая подложка только на subheading-row (HEX F2F2F2).
        _set_cell_shading(sub_cell, "F2F2F2")
        # Subheading не должна оставаться висячей в конце страницы:
        # «keepNext» прикрепляет её к следующей (первой data-row группы).
        sub_para.paragraph_format.keep_with_next = True
        _set_row_keep_together(sub_row, allow_split=False)

        for a in group:
            num += 1
            data_row = tbl.add_row()
            ident = a.cad_number or a.inv_number or "—"
            name_text = a.name or "—"
            addr_text = a.address or "—"
            benef_text = beneficiary_cell(a)
            data_row.cells[0].text = str(num)
            data_row.cells[1].text = ident
            data_row.cells[2].text = name_text
            data_row.cells[3].text = addr_text
            data_row.cells[4].text = benef_text
            # Колонка «№»: содержимое по центру, явный bold=False.
            num_cell = data_row.cells[0]
            num_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            num_para = num_cell.paragraphs[0]
            num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in num_para.runs:
                r.bold = False
            # Висячие строки: запрещаем разрыв строки между страницами,
            # ЕСЛИ суммарная высота строки укладывается в ~30% страницы.
            # Эвристика по длине текста: Arial 10pt в колонках 5.5/3.5 см
            # → ~34/22 chars/line; 16 строк (≈30% A4) → 544/350 chars.
            allow_split = (
                len(name_text) > 544
                or len(addr_text) > 350
                or len(benef_text) > 350
            )
            _set_row_keep_together(data_row, allow_split=allow_split)

    _set_column_widths(tbl, widths_cm)


def _set_cell_shading(cell, fill_hex: str) -> None:
    """Заливка ячейки сплошным цветом (hex без #). Используется только на
    subheading-row группы — отметить визуально, не контрастно."""
    from docx.oxml.ns import qn  # type: ignore[import]
    from docx.oxml import OxmlElement  # type: ignore[import]
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _set_row_keep_together(row, allow_split: bool) -> None:
    """Управляет переносом строки таблицы между страницами.

    `allow_split=False` → строка ОБЯЗАНА остаться на одной странице
    (`w:cantSplit`). Если содержимое не помещается на странице, Word
    переносит всю строку целиком на следующую.

    `allow_split=True` → разрешаем разрыв (для очень длинных значений,
    которые сами по себе занимают больше 30% высоты страницы — иначе
    cantSplit мог бы вызвать клиппинг).
    """
    from docx.oxml.ns import qn  # type: ignore[import]
    from docx.oxml import OxmlElement  # type: ignore[import]
    tr_pr = row._tr.get_or_add_trPr()
    for existing in tr_pr.findall(qn("w:cantSplit")):
        tr_pr.remove(existing)
    if not allow_split:
        cant = OxmlElement("w:cantSplit")
        tr_pr.append(cant)


def _mark_row_as_header(row) -> None:
    """Помечает строку как заголовок таблицы — Word повторит её на каждой
    следующей странице, если таблица переносится."""
    from docx.oxml.ns import qn  # type: ignore[import]
    from docx.oxml import OxmlElement  # type: ignore[import]
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_column_widths(tbl, widths_cm: tuple[float, ...]) -> None:
    """Жёстко задаём ширины колонок (python-docx требует прописать каждой
    ячейке и колонку, и `tblGrid` через OXML, иначе Word проигнорирует)."""
    from docx.shared import Cm  # type: ignore[import]
    from docx.oxml.ns import qn  # type: ignore[import]
    from docx.oxml import OxmlElement  # type: ignore[import]

    # Обновляем tblGrid в OXML
    tbl_el = tbl._tbl
    grid = tbl_el.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl_el.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w_cm in widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(Cm(w_cm).emu / 635)))  # EMU → twips
        grid.append(gc)

    # Прописываем ширину каждой ячейке (иначе Word подгонит автоматически).
    # Строки с объединёнными ячейками (subheading-row) пропускаем — для них
    # ширина = сумме колонок и задаётся неявно через tblGrid.
    for row in tbl.rows:
        cells = row.cells
        unique_tcs = {id(c._tc) for c in cells}
        if len(unique_tcs) <= 1:
            continue
        for i, cell in enumerate(cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


# ─── main ───────────────────────────────────────────────────────────────────

BANNER = "═" * 70


def main() -> int:
    cp(f"\n{BANNER}", C.B)
    cp("  ekcelo · ТЗ-1: Список объектов для исследования", C.B)
    cp(f"{BANNER}\n", C.B)

    project = ask_path("Введите путь к папке проекта")
    if not project.is_dir():
        cp(f"✗ Это не папка: {project}", C.R)
        return 1

    out_dir = ask_path(
        "Введите путь к папке для сохранения DOCX",
        must_exist=False, default=str(project / "reports"),
    )
    out_dir.mkdir(parents=True, exist_ok=True)

    acc = Accumulator()
    cp(f"\n· Папка проекта: {project}", C.CY)
    cp(f"· Папка выгрузки: {out_dir}", C.CY)

    cp("\n[1/3] Гибридная загрузка из _data/structure.json …", C.B)
    used_snapshot = load_from_structure(project, acc)
    if not used_snapshot:
        cp("  ℹ Snapshot не найден или пуст — будет только сырое сканирование.",
           C.Y)

    do_raw = True
    if used_snapshot:
        do_raw = ask_yn(
            "Дополнить сырым сканированием xlsx/pdf/xml в проекте?",
            default=True,
        )

    if do_raw:
        cp("\n[2/3] Рекурсивное сканирование сырых файлов …", C.B)
        n_pdf, n_xml, n_xlsx = scan_raw(project, acc)
        cp(f"  ✓ Просмотрено: PDF={n_pdf}, XML={n_xml}, XLSX={n_xlsx}", C.G)
    else:
        cp("\n[2/3] Сырое сканирование пропущено.", C.CY)

    cp("\n[3/3] Связывание активов с бенефициарами …", C.B)
    link_holders(acc)

    all_assets = list(acc.assets.values())
    linked = [a for a in all_assets if a.has_link()]
    cp(f"  Всего объектов: {len(all_assets)}; со связью: {len(linked)}", C.CY)
    cp(f"  Бенефициаров найдено: {len(acc.beneficiaries)}", C.CY)

    if not all_assets:
        cp("\n✗ Не найдено ни одного объекта. ТЗ-1 не будет сформировано.", C.R)
        return 1

    now = datetime.now()
    ts = now.strftime("%Y%m%d_%H%M%S")
    out_docx = out_dir / f"TZ1{ts}.docx"
    out_md = out_dir / f"TZ1{ts}.md"
    meta_dict = {
        "time": now.strftime("%H:%M:%S"),
        "date": now.strftime("%d.%m.%Yг."),
        "doc_num": now.strftime("%Y%m%d"),
        "doc_date": now.strftime("%d.%m.%Yг."),
        "root": str(project),
        "total": len(all_assets),
        "linked": len(linked),
        "beneficiaries": len(acc.beneficiaries),
    }

    # MD-выгрузка — без зависимостей (text only), идёт первой, чтобы при
    # отсутствии python-docx у пользователя был хотя бы Markdown.
    build_md(table1=linked, out_path=out_md, meta=meta_dict)
    cp(f"\n✓ Готово: {out_md}", C.G)

    try:
        build_docx(table1=linked, out_path=out_docx, meta=meta_dict)
        cp(f"✓ Готово: {out_docx}", C.G)
    except ImportError as e:
        cp(f"\n⚠ python-docx не установлен: {e}", C.Y)
        cp("  Установите: pip install python-docx openpyxl pdfplumber", C.Y)
        cp(f"  DOCX пропущен; MD доступен: {out_md}", C.Y)
    return 0


if __name__ == "__main__":
    try:
        sys.exit(main())
    except KeyboardInterrupt:
        cp("\n✗ Прервано пользователем.", C.Y)
        sys.exit(130)
