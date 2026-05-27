#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""pirushin_sosn_rocha_13_assemble_contract_v1.py

ekcelo · Assembler договоров (Surveycontract).

Читает MD-output скриптов 10/11/12 из `<project>/Surveycontract/{tz1-content,
body,tz2-calculation}/`, интерактивно компонует выбранные компоненты в
финальную сборку (`gotovo/`) и сохраняет конфигурацию в `sborki/`.

Idempotent re-build: при наличии существующего `sborki/<name>.json`
скрипт предлагает «пересобрать с теми же компонентами» (без re-prompt'а).

Версионирование: поле `_parent_sborka` указывает на predшествующую
версию того же договора (для допсоглашений). Поле `_parent_contract` —
на parent-номер договора (для субподрядов).
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from datetime import datetime
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))


# ─── ANSI ─────────────────────────────────────────────────────────────────


class C:
    R = "\033[31m"
    G = "\033[32m"
    Y = "\033[33m"
    B = "\033[36m"
    O = "\033[0m"  # noqa: E741


def cp(msg: str, col: str = "") -> None:
    print(f"{col}{msg}{C.O}" if col else msg)


# ─── Сканирование компонентов ──────────────────────────────────────────────


GROUPS = {
    "tz1": "tz1-content",
    "body": "body",
    "calc": "tz2-calculation",
}

# Имена выходных файлов 10/11/12 — извлекаем datetime + (опц.) номер договора.
_RX_TZ1 = re.compile(r"^TZ1[_-]?(\d{8}[_-]?\d{6})\.md$")
_RX_BODY = re.compile(r"^Contract_([^_]+)_(\d{8}[_-]?\d{6})\.md$")
_RX_CALC = re.compile(r"^Appendix2_([^_]+)_(\d{8}[_-]?\d{6})\.md$")


def _meta_from_head(md_path: Path) -> dict:
    """Извлекает predmet_kind / summary из первых строк MD."""
    meta = {"predmet_kind": None, "summary": None}
    try:
        head = md_path.read_text(encoding="utf-8", errors="ignore").splitlines()[:20]
    except OSError:
        return meta
    blob = "\n".join(head)
    if "ГК-39" in blob:
        meta["predmet_kind"] = "gk39"
    elif "135-ФЗ" in blob:
        meta["predmet_kind"] = "fz135"
    # Summary = первая непустая строка после заголовка
    for line in head[1:]:
        s = line.strip()
        if s and not s.startswith("#"):
            meta["summary"] = s[:80]
            break
    return meta


def scan_group(sc_dir: Path, group: str) -> list[dict]:
    """Возвращает [{path, name, ts, number, predmet_kind, summary}, …]."""
    g_dir = sc_dir / GROUPS[group]
    if not g_dir.is_dir():
        return []
    items: list[dict] = []
    for md in sorted(g_dir.glob("*.md")):
        ts = number = None
        if group == "tz1":
            if m := _RX_TZ1.match(md.name):
                ts = m.group(1)
        elif group == "body":
            if m := _RX_BODY.match(md.name):
                number, ts = m.group(1), m.group(2)
        elif group == "calc":
            if m := _RX_CALC.match(md.name):
                number, ts = m.group(1), m.group(2)
        meta = _meta_from_head(md)
        items.append({
            "path": md,
            "name": md.name,
            "ts": ts,
            "number": number,
            "predmet_kind": meta["predmet_kind"],
            "summary": meta["summary"],
        })
    return items


# ─── Интерактивный выбор ──────────────────────────────────────────────────


def _pick_one(group_label: str, items: list[dict], auto_latest: bool = False) -> dict | None:
    if not items:
        cp(f"  ⚠ В группе [{group_label}] нет MD-файлов.", C.Y)
        return None
    cp(f"\nГруппа [{group_label}]:", C.B)
    for i, it in enumerate(items, 1):
        suffix = []
        if it["number"]:
            suffix.append(f"№{it['number']}")
        if it["predmet_kind"]:
            suffix.append(it["predmet_kind"])
        if it["summary"]:
            suffix.append(it["summary"])
        tail = "  (" + " / ".join(suffix) + ")" if suffix else ""
        marker = " ← последний" if i == len(items) else ""
        cp(f"  {i}. {it['name']}{tail}{marker}", C.B)
    if auto_latest:
        cp(f"    [auto: {items[-1]['name']}]", C.G)
        return items[-1]
    while True:
        raw = input("  Выбор (Enter — последний / 0 — пропустить): ").strip()
        if raw == "":
            return items[-1]
        if raw == "0":
            return None
        if raw.isdigit() and 1 <= int(raw) <= len(items):
            return items[int(raw) - 1]
        cp(f"  ⚠ Ожидается 1..{len(items)} / Enter / 0.", C.Y)


def _pick_upd(sc_dir: Path, auto: bool = False) -> Path | None:
    upd_dir = sc_dir / "upd"
    xmls = sorted(upd_dir.glob("*.xml")) if upd_dir.is_dir() else []
    if auto:
        return xmls[-1] if xmls else None
    cp("\nГруппа [upd] (опционально):", C.B)
    if not xmls:
        cp("  ⚠ В upd/ нет XML-файлов. Пропускаем.", C.Y)
        return None
    for i, p in enumerate(xmls, 1):
        cp(f"  {i}. {p.name}", C.B)
    while True:
        raw = input("  Выбор (Enter — без УПД / номер): ").strip()
        if raw == "":
            return None
        if raw.isdigit() and 1 <= int(raw) <= len(xmls):
            return xmls[int(raw) - 1]
        cp("  ⚠ Ожидается номер или Enter.", C.Y)


# ─── Сохранение конфигурации + сборка ──────────────────────────────────────


def assemble(
    project: Path,
    *,
    tz1: Path,
    body: Path,
    calc: Path,
    upd: Path | None,
    contract_number: str | None = None,
    contract_date: str | None = None,
    version: int = 1,
    predmet_kind: str | None = None,
    parent_sborka: str | None = None,
    parent_contract: str | None = None,
    formats: tuple[str, ...] = ("md", "json"),
) -> dict:
    """Сохраняет sborka-конфиг + финальный артефакт в gotovo/.

    Возвращает {sborka_path, gotovo_paths: dict[ext->path]}.
    """
    sc = project / "Surveycontract"
    sborki_dir = sc / "sborki"
    gotovo_dir = sc / "gotovo"
    sborki_dir.mkdir(parents=True, exist_ok=True)
    gotovo_dir.mkdir(parents=True, exist_ok=True)

    now = datetime.now()
    ts = now.strftime("%Y%m%d_%H%M%S")
    contract_number = contract_number or now.strftime("%Y%m%d") + "-1"
    contract_date = contract_date or now.strftime("%d.%m.%Yг.")
    safe_num = contract_number.replace("/", "-").replace(" ", "_")

    config = {
        "contract_number": contract_number,
        "contract_date": contract_date,
        "version": version,
        "predmet_kind": predmet_kind,
        "components": {
            "tz1": str(tz1.relative_to(sc)) if tz1 else None,
            "body": str(body.relative_to(sc)) if body else None,
            "calc": str(calc.relative_to(sc)) if calc else None,
            "upd": str(upd.relative_to(sc)) if upd else None,
        },
        "_parent_sborka": parent_sborka,
        "_parent_contract": parent_contract,
        "_created_at": now.isoformat(timespec="seconds"),
    }

    sborka_name = f"contract_{safe_num}_v{version}_{ts}.json"
    sborka_path = sborki_dir / sborka_name
    sborka_path.write_text(
        json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    # Финальная сборка — concat MD-секций с заголовками.
    parts = []
    if tz1 and tz1.exists():
        parts.append("# Приложение №1 — Техническое задание №1\n\n" +
                     tz1.read_text(encoding="utf-8"))
    if body and body.exists():
        parts.append("# Тело договора\n\n" +
                     body.read_text(encoding="utf-8"))
    if calc and calc.exists():
        parts.append("# Приложение №2 — Календарный план и калькуляция\n\n" +
                     calc.read_text(encoding="utf-8"))
    if upd and upd.exists():
        parts.append("# УПД (приложен)\n\n```xml\n" +
                     upd.read_text(encoding="utf-8") +
                     "\n```")
    combined_md = "\n\n---\n\n".join(parts)

    base = f"{safe_num}_{contract_date.replace('.', '-').rstrip('г-')}_v{version}_{ts}"
    paths: dict[str, str] = {}
    if "md" in formats:
        p = gotovo_dir / f"{base}.md"
        p.write_text(combined_md, encoding="utf-8")
        paths["md"] = str(p)
    if "json" in formats:
        p = gotovo_dir / f"{base}.json"
        artifact = dict(config)
        artifact["combined_md"] = combined_md
        p.write_text(json.dumps(artifact, ensure_ascii=False, indent=2),
                     encoding="utf-8")
        paths["json"] = str(p)
    if "docx" in formats:
        try:
            from docx import Document  # type: ignore

            doc = Document()
            for section in combined_md.split("\n\n---\n\n"):
                first_line = section.split("\n", 1)[0].lstrip("# ").strip()
                doc.add_heading(first_line or "Раздел", level=1)
                body_md = section.split("\n", 1)[1] if "\n" in section else ""
                for para in body_md.split("\n"):
                    if para.strip():
                        doc.add_paragraph(para)
            p = gotovo_dir / f"{base}.docx"
            doc.save(str(p))
            paths["docx"] = str(p)
        except ImportError:
            cp("  ⚠ python-docx не установлен — DOCX пропущен.", C.Y)

    return {"sborka_path": str(sborka_path), "gotovo_paths": paths,
            "config": config}


# ─── CLI ──────────────────────────────────────────────────────────────────


def _ask(prompt: str, default: str = "") -> str:
    suffix = f" [{default}]" if default else ""
    raw = input(f"{prompt}{suffix}: ").strip()
    return raw or default


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(
        description="ekcelo · assembler договоров (Surveycontract)",
    )
    ap.add_argument("--project", type=Path, required=True,
                    help="папка проекта (содержит Surveycontract/)")
    ap.add_argument("--auto-latest", action="store_true",
                    help="non-interactive: выбирать последний файл в каждой группе")
    ap.add_argument("--formats", default="md,json",
                    help="форматы финальной сборки через запятую (md,json,docx)")
    args = ap.parse_args(argv)

    project = args.project.expanduser()
    sc = project / "Surveycontract"
    if not sc.is_dir():
        cp(f"✗ Не найдено: {sc}. Сначала запустите 07_init_project_v3.", C.R)
        return 1

    tz1_items = scan_group(sc, "tz1")
    body_items = scan_group(sc, "body")
    calc_items = scan_group(sc, "calc")

    tz1 = _pick_one("tz1-content", tz1_items, auto_latest=args.auto_latest)
    body = _pick_one("body", body_items, auto_latest=args.auto_latest)
    calc = _pick_one("tz2-calculation", calc_items, auto_latest=args.auto_latest)
    upd = _pick_upd(sc, auto=args.auto_latest)

    if not (tz1 and body and calc):
        cp("✗ Нужны все 3 компонента (tz1, body, calc). Прервано.", C.R)
        return 1

    if args.auto_latest:
        number = body["number"] or tz1["number"] or "AUTO"
        kind = body["predmet_kind"] or tz1["predmet_kind"]
    else:
        number = _ask("Номер договора", body["number"] or "")
        kind_default = body["predmet_kind"] or tz1["predmet_kind"] or ""
        kind = _ask("Тип договора (gk39/fz135)", kind_default)

    formats = tuple(f.strip() for f in args.formats.split(",") if f.strip())

    result = assemble(
        project,
        tz1=tz1["path"],
        body=body["path"],
        calc=calc["path"],
        upd=upd,
        contract_number=number or None,
        predmet_kind=kind or None,
        formats=formats,
    )
    cp("\n✓ Готово:", C.G)
    cp(f"    sborka:  {result['sborka_path']}", C.G)
    for ext, p in result["gotovo_paths"].items():
        cp(f"    gotovo {ext}: {p}", C.G)
    return 0


if __name__ == "__main__":
    try:
        sys.exit(main())
    except KeyboardInterrupt:
        cp("\n✗ Прервано пользователем.", C.Y)
        sys.exit(130)
