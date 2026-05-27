#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""ekcelo · 11_make_contract_body_v1 — Тело рамочного договора (DOCX + MD).

Формирует «Договор № …» возмездного оказания услуг по сбору и
систематизации данных для оценки. Структура DOCX:
  • шапка / Стороны;
  • 1. Предмет (вариант: «ГК РФ гл. 39» либо «бесшовный маппинг
    под отчёт по 135-ФЗ»);
  • 2. Цена и порядок расчётов;
  • 3. Сроки;
  • 4. Права и обязанности;
  • 5. Конфиденциальность;
  • 6. Ответственность;
  • 7. Разрешение споров;
  • 8. Заключительные положения;
  • Перечень приложений (обязательные);
  • 9. Реквизиты сторон (2-кол. таблица без рамок);
  • подписи.

Приложения 1 и 2 формируются отдельными скриптами:
  • 10_make_tz1_for_contract_ekcelo_v1.py  — Приложение №1 (ТЗ-1);
  • 12_make_contract_appendix2_v1.py        — Приложение №2 (план/услуги).

Реквизиты сторон читаются из parties.json (см. obsidian/Database/
parties.example.json). Если JSON не найден — все поля вводятся вручную.

CLI:
  python pirushin_sosn_rocha_11_make_contract_body_v1.py
"""

from __future__ import annotations

import json
import sys
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

# ───────────────────────────── CLI helpers ──────────────────────────────

BANNER = "═" * 66


class C:
    R = "\033[31m"; G = "\033[32m"; Y = "\033[33m"; B = "\033[1m"
    CY = "\033[36m"; X = "\033[0m"


def cp(t: str = "", c: str = "") -> None:
    print(f"{c}{t}{C.X}" if c else t)


def ask(prompt: str, default: str = "") -> str:
    suffix = f" [{default}]" if default else ""
    s = input(f"{prompt}{suffix}: ").strip()
    return s or default


def ask_path(prompt: str, must_exist: bool = True,
             default: str = "") -> Path:
    while True:
        raw = ask(prompt, default)
        if not raw:
            cp("  ⚠ Путь не задан.", C.Y); continue
        p = Path(raw.strip().strip('"').strip("'")).expanduser()
        if must_exist and not p.exists():
            cp(f"  ⚠ Не найдено: {p}", C.Y); continue
        return p


def ask_int(prompt: str, default: int, lo: int, hi: int) -> int:
    while True:
        raw = ask(prompt, str(default))
        try:
            v = int(raw)
        except ValueError:
            cp(f"  ⚠ Нужно число от {lo} до {hi}.", C.Y); continue
        if lo <= v <= hi:
            return v
        cp(f"  ⚠ Вне диапазона {lo}…{hi}.", C.Y)


def ask_choice(prompt: str, options: list[str], default: int = 1) -> int:
    """Выбор из пронумерованного списка. Возвращает 1-based индекс."""
    cp(prompt, C.B)
    for i, lbl in enumerate(options, 1):
        marker = " ←" if i == default else ""
        cp(f"  {i}. {lbl}{marker}", C.CY)
    return ask_int("Выберите номер", default, 1, len(options))


# ───────────────────────────── parties.json ─────────────────────────────

@dataclass
class Party:
    """Сводные реквизиты одной стороны (Исполнителя или Заказчика)."""

    short_name: str = ""
    full_name: str = ""
    inn: str = ""
    kpp: str = ""        # для юрлиц
    ogrn: str = ""       # ОГРН или ОГРНИП
    address: str = ""
    bank: str = ""
    account: str = ""
    bik: str = ""
    ks: str = ""
    phone: str = ""
    email: str = ""
    signatory_name: str = ""
    signatory_position: str = ""
    signatory_basis: str = ""

    @classmethod
    def from_dict(cls, d: dict[str, Any]) -> "Party":
        """Берёт известные ключи; «ogrnip» допустим как алиас «ogrn»."""
        return cls(
            short_name=str(d.get("short_name", "") or ""),
            full_name=str(d.get("full_name", "") or ""),
            inn=str(d.get("inn", "") or ""),
            kpp=str(d.get("kpp", "") or ""),
            ogrn=str(d.get("ogrn", "") or d.get("ogrnip", "") or ""),
            address=str(d.get("address", "") or ""),
            bank=str(d.get("bank", "") or ""),
            account=str(d.get("account", "") or ""),
            bik=str(d.get("bik", "") or ""),
            ks=str(d.get("ks", "") or ""),
            phone=str(d.get("phone", "") or ""),
            email=str(d.get("email", "") or ""),
            signatory_name=str(d.get("signatory_name", "") or ""),
            signatory_position=str(d.get("signatory_position", "") or ""),
            signatory_basis=str(d.get("signatory_basis", "") or ""),
        )


def _candidate_parties_paths(project: Path | None) -> list[Path]:
    """Места, где скрипт ищет parties.json по приоритету."""
    here = Path(__file__).resolve().parent
    repo = here.parent.parent  # parser/scripts/ → parser/ → repo
    candidates = []
    if project:
        candidates.append(project / "_data" / "parties.json")
    candidates.extend([
        Path.cwd() / "parties.json",
        repo / "parser" / "_data" / "parties.json",
        repo / "obsidian" / "Database" / "parties.json",
    ])
    return candidates


def load_parties_file(project: Path | None) -> tuple[
        list[Party], list[Party], Path | None]:
    """Возвращает (executors, clients, source_path|None).

    Если JSON не найден — пустые списки и None.
    """
    for p in _candidate_parties_paths(project):
        if p.is_file():
            try:
                data = json.loads(p.read_text(encoding="utf-8"))
            except (OSError, json.JSONDecodeError) as e:
                cp(f"  ⚠ Не удалось прочитать {p}: {e}", C.Y)
                continue
            execs = [Party.from_dict(x) for x in data.get("executors", [])]
            clis = [Party.from_dict(x) for x in data.get("clients", [])]
            return execs, clis, p
    return [], [], None


def pick_or_enter_party(
        role: str, options: list[Party]) -> Party:
    """Меню: выбрать из JSON, ввести руками, либо «оставить пусто»."""
    cp(f"\n— {role}: реквизиты —", C.B)
    labels: list[str] = []
    for p in options:
        labels.append(
            f"{p.short_name or '(без названия)'} "
            f"(ИНН {p.inn or '___'})"
        )
    labels.append("Ввести вручную")
    labels.append("Оставить плейсхолдеры «___» (заполню в Word)")

    default_idx = 1 if options else len(labels) - 1
    choice = ask_choice("Выберите вариант:", labels, default=default_idx)

    if choice <= len(options):
        chosen = options[choice - 1]
        cp(f"  ✓ Выбран: {chosen.short_name}", C.G)
        return chosen
    if choice == len(labels) - 1:
        return enter_party_manually(role)
    # «Оставить пусто»
    cp("  ℹ Будут поставлены плейсхолдеры — заполните в Word.", C.CY)
    return Party()


def enter_party_manually(role: str) -> Party:
    """Пошаговый ввод реквизитов с консоли. Пустое поле → '' (плейсхолдер)."""
    cp(f"\n  Ввод реквизитов «{role}» (Enter = пусто/«___»):", C.CY)
    return Party(
        short_name=ask("  Краткое наименование (напр. ООО «Пример»)"),
        full_name=ask("  Полное наименование"),
        inn=ask("  ИНН"),
        kpp=ask("  КПП (для юрлиц; для ИП оставьте пустым)"),
        ogrn=ask("  ОГРН/ОГРНИП"),
        address=ask("  Юридический адрес"),
        bank=ask("  Банк"),
        account=ask("  Расчётный счёт"),
        bik=ask("  БИК"),
        ks=ask("  Корр. счёт"),
        phone=ask("  Телефон"),
        email=ask("  Email"),
        signatory_name=ask("  Подписант (ФИО)"),
        signatory_position=ask("  Должность подписанта"),
        signatory_basis=ask("  Действует на основании (Устава / доверенности №…)"),
    )


# ───────────────────── контракт: метаданные & текст ─────────────────────

# Два варианта формулировок Предмета и Заключительных положений.
CONTRACT_KIND_GK = "gk"     # ст. 779-783 ГК РФ
CONTRACT_KIND_FZ135 = "fz135"  # бесшовный маппинг под отчёт 135-ФЗ


@dataclass
class ContractMeta:
    number: str
    date_str: str        # «27 мая 2026г.»
    city: str
    kind: str            # CONTRACT_KIND_GK | CONTRACT_KIND_FZ135
    executor: Party
    client: Party

    def title(self) -> str:
        return f"Договор № {self.number} возмездного оказания услуг"


_RU_MONTHS = [
    "", "января", "февраля", "марта", "апреля", "мая", "июня",
    "июля", "августа", "сентября", "октября", "ноября", "декабря",
]


def ru_date(dt: datetime) -> str:
    """27 мая 2026г."""
    return f"{dt.day} {_RU_MONTHS[dt.month]} {dt.year}г."


def party_header_line(p: Party, role: str) -> str:
    """«ООО «Пример» (ИНН …, ОГРН …), в лице … действующего на основании …,
    далее — Заказчик» — одна-две строки для шапки договора."""
    full = p.full_name or p.short_name or f"___ ({role}) ___"
    inn = p.inn or "___"
    ogrn = p.ogrn or "___"
    sig_name = p.signatory_name or "___"
    sig_pos = p.signatory_position or "___"
    basis = p.signatory_basis or "Устава"
    return (
        f"{full} (ИНН {inn}, ОГРН {ogrn}), в лице "
        f"{sig_pos} {sig_name}, действующего(-ей) на основании {basis} "
        f"(далее — {role})"
    )


# Тексты пунктов. Хранятся как (заголовок, абзацы). Абзацы рендерятся
# через add_paragraph; пустые строки — отступ. Используем плейсхолдеры
# {placeholders} для конкретных полей.

def _section_predmet(kind: str) -> tuple[str, list[str]]:
    common = [
        "1.1. Исполнитель обязуется по заданию Заказчика оказать услуги "
        "по сбору, систематизации и идентификации данных об объектах "
        "недвижимости и оборудовании Заказчика, перечень которых указан "
        "в Приложении №1 (обязательном) — Техническом задании №1 — к "
        "настоящему Договору, а Заказчик обязуется принять и оплатить "
        "эти услуги.",
        "1.2. Состав работ (услуг), сроки выполнения и стоимость "
        "установлены в Приложении №2 (обязательном) — Календарном плане, "
        "Перечне источников информации и Чеклисте услуг — к настоящему "
        "Договору.",
    ]
    if kind == CONTRACT_KIND_FZ135:
        common.append(
            "1.3. Результаты Работ являются исходными данными, "
            "пригодными для бесшовного маппинга в Отчёт об оценке, "
            "формируемый в рамках отдельного договора об оценке "
            "(заключаемого в соответствии с Федеральным законом от "
            "29.07.1998 № 135-ФЗ «Об оценочной деятельности в "
            "Российской Федерации»). Настоящий Договор НЕ является "
            "договором об оценочной деятельности в смысле ст. 9 135-ФЗ."
        )
    else:
        common.append(
            "1.3. Настоящий Договор является договором возмездного "
            "оказания услуг в смысле главы 39 (ст. 779-783) "
            "Гражданского кодекса Российской Федерации. Результаты Работ "
            "оформляются актом сдачи-приёмки оказанных услуг."
        )
    return "1. ПРЕДМЕТ ДОГОВОРА", common


def _section_price() -> tuple[str, list[str]]:
    return "2. ЦЕНА И ПОРЯДОК РАСЧЁТОВ", [
        "2.1. Общая стоимость работ по настоящему Договору указана "
        "итоговой строкой «ИТОГО» в Приложении №2.",
        "2.2. НДС: указывается в соответствии с применяемым налоговым "
        "режимом Исполнителя (УСН/НПД — не облагается; ОСН — 20 %).",
        "2.3. Оплата производится в течение 5 (пяти) рабочих дней с "
        "даты подписания Сторонами акта сдачи-приёмки оказанных услуг, "
        "путём перечисления денежных средств на расчётный счёт "
        "Исполнителя, указанный в разделе 9 настоящего Договора.",
        "2.4. По соглашению Сторон может быть предусмотрен аванс "
        "(не более 50 % от суммы Приложения №2) — оформляется "
        "дополнительным соглашением.",
    ]


def _section_terms() -> tuple[str, list[str]]:
    return "3. СРОКИ ВЫПОЛНЕНИЯ", [
        "3.1. Сроки выполнения отдельных этапов работ установлены "
        "Календарным планом (Приложение №2).",
        "3.2. Настоящий Договор вступает в силу с момента подписания "
        "обеими Сторонами и действует до полного исполнения "
        "Сторонами своих обязательств.",
        "3.3. Сроки могут быть продлены по уважительным причинам "
        "(простой по вине Заказчика, ожидание согласования выезда, "
        "ожидание дополнительных материалов) — соразмерно фактической "
        "задержке. Простой оплачивается отдельной строкой «Простой при "
        "ожидании…» (Приложение №2).",
    ]


def _section_rights() -> tuple[str, list[str]]:
    return "4. ПРАВА И ОБЯЗАННОСТИ СТОРОН", [
        "4.1. Исполнитель обязан:",
        "  4.1.1. оказать услуги в установленные Приложением №2 сроки;",
        "  4.1.2. сохранять конфиденциальность сведений, полученных "
        "в ходе исполнения Договора;",
        "  4.1.3. передать Заказчику результаты Работ в формате DOCX/PDF/"
        "KMZ-KML и (при наличии) doc-фотофиксацию и граф связей;",
        "  4.1.4. размещать блок фотографий на файлообменнике и "
        "предоставлять ссылку на скачивание.",
        "4.2. Заказчик обязан:",
        "  4.2.1. передать Исполнителю необходимую исходную документацию "
        "(перечень — в Приложении №2);",
        "  4.2.2. обеспечить согласование и проведение выезда "
        "представителя Исполнителя к местоположению объектов "
        "идентификации с представителем балансодержателя;",
        "  4.2.3. возместить транспортные расходы и расходы на "
        "проживание (вне точки дислокации) специалистов Исполнителя "
        "по фактическим затратам — отдельной строкой Приложения №2;",
        "  4.2.4. оплатить оказанные услуги в порядке раздела 2.",
    ]


def _section_confid() -> tuple[str, list[str]]:
    return "5. КОНФИДЕНЦИАЛЬНОСТЬ", [
        "5.1. Стороны обязуются не разглашать третьим лицам сведения, "
        "ставшие известными им в ходе исполнения настоящего Договора, "
        "за исключением случаев, прямо предусмотренных "
        "законодательством Российской Федерации.",
        "5.2. Обязательство о конфиденциальности действует в течение "
        "3 (трёх) лет с момента прекращения настоящего Договора.",
    ]


def _section_liability() -> tuple[str, list[str]]:
    return "6. ОТВЕТСТВЕННОСТЬ СТОРОН", [
        "6.1. Стороны несут ответственность за неисполнение или "
        "ненадлежащее исполнение обязательств в соответствии с "
        "действующим законодательством Российской Федерации.",
        "6.2. За просрочку оплаты Заказчик уплачивает Исполнителю "
        "пеню в размере 0,1 % от просроченной суммы за каждый день "
        "просрочки, но не более 10 % от общей стоимости работ.",
    ]


def _section_disputes() -> tuple[str, list[str]]:
    return "7. РАЗРЕШЕНИЕ СПОРОВ", [
        "7.1. Споры, возникающие из настоящего Договора, разрешаются "
        "Сторонами путём переговоров.",
        "7.2. При недостижении согласия — в арбитражном суде по месту "
        "нахождения Исполнителя.",
    ]


def _section_final(kind: str) -> tuple[str, list[str]]:
    base = [
        "8.1. Настоящий Договор регулируется Гражданским кодексом "
        "Российской Федерации (глава 39) и Налоговым кодексом "
        "Российской Федерации.",
        "8.2. Все изменения и дополнения к настоящему Договору "
        "оформляются дополнительными соглашениями в письменной форме, "
        "подписываемыми обеими Сторонами.",
        "8.3. Настоящий Договор составлен в 2 (двух) экземплярах, "
        "имеющих равную юридическую силу, по одному для каждой Стороны.",
    ]
    if kind == CONTRACT_KIND_FZ135:
        base.insert(0,
            "8.0. Стороны подтверждают, что результаты Работ "
            "предназначены для использования в отчёте об оценке, "
            "формируемом в порядке Федерального закона от 29.07.1998 "
            "№ 135-ФЗ; конкретные требования к составу Отчёта "
            "согласовываются отдельным договором об оценке."
        )
    return "8. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ", base


def _section_appendices() -> tuple[str, list[str]]:
    return "ПРИЛОЖЕНИЯ К ДОГОВОРУ (обязательные)", [
        "№1. Техническое задание №1 — Список объектов для исследования.",
        "№2. Календарный план, Перечень источников информации, "
        "Чеклист услуг со стоимостью.",
    ]


# ───────────────────────────── DOCX builder ─────────────────────────────

def build_docx(meta: ContractMeta, out_path: Path) -> Path:
    from docx import Document  # type: ignore[import]
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for sec in doc.sections:
        sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
        sec.top_margin = sec.bottom_margin = Cm(2.0)
        sec.left_margin = sec.right_margin = Cm(2.0)
    st = doc.styles["Normal"]
    st.font.name = "Arial"; st.font.size = Pt(11)

    # ── Шапка: «Договор № …» по центру жирным.
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(meta.title())
    run.bold = True; run.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("по сбору и систематизации данных для целей оценки")

    # ── Город / Дата (по краям).
    cd = doc.add_paragraph()
    tab = cd.paragraph_format
    cd.add_run(f"г. {meta.city or '___'}")
    cd.add_run("\t" * 8)
    cd.add_run(meta.date_str)

    # ── Преамбула со Сторонами.
    doc.add_paragraph()
    p_exec = doc.add_paragraph()
    p_exec.add_run(party_header_line(meta.executor, "Исполнитель") + ",")
    p_exec.add_run(" с одной стороны, и")

    p_client = doc.add_paragraph()
    p_client.add_run(party_header_line(meta.client, "Заказчик") + ",")
    p_client.add_run(" с другой стороны,")

    p_join = doc.add_paragraph()
    p_join.add_run("далее совместно именуемые «Стороны», ")
    p_join.add_run("заключили настоящий Договор о нижеследующем:")

    # ── Разделы 1-8 + перечень приложений.
    sections = [
        _section_predmet(meta.kind),
        _section_price(),
        _section_terms(),
        _section_rights(),
        _section_confid(),
        _section_liability(),
        _section_disputes(),
        _section_final(meta.kind),
        _section_appendices(),
    ]
    for heading, paragraphs in sections:
        doc.add_paragraph()
        h = doc.add_paragraph()
        r = h.add_run(heading); r.bold = True
        for txt in paragraphs:
            doc.add_paragraph(txt)

    # ── Раздел 9. Реквизиты сторон (2-кол. таблица без рамок).
    doc.add_paragraph()
    h9 = doc.add_paragraph()
    r9 = h9.add_run("9. РЕКВИЗИТЫ СТОРОН"); r9.bold = True

    _emit_requisites_table(doc, meta.executor, meta.client)

    # ── Подписи (под таблицей реквизитов, тоже 2 кол.)
    doc.add_paragraph()
    _emit_signatures_block(doc, meta.executor, meta.client)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def _requisites_lines(p: Party, role: str) -> list[str]:
    """Список строк блока реквизитов одной стороны."""
    short = p.short_name or f"___ ({role}) ___"
    return [
        f"{role}:",
        short,
        f"ИНН: {p.inn or '___'}",
        f"КПП: {p.kpp or '—'}",
        f"ОГРН/ОГРНИП: {p.ogrn or '___'}",
        f"Адрес: {p.address or '___'}",
        f"Банк: {p.bank or '___'}",
        f"Р/с: {p.account or '___'}",
        f"БИК: {p.bik or '___'}",
        f"К/с: {p.ks or '___'}",
        f"Тел.: {p.phone or '___'}",
        f"Email: {p.email or '___'}",
    ]


def _emit_requisites_table(doc, exe: Party, cli: Party) -> None:
    """2-колонка без рамок: слева Исполнитель, справа Заказчик."""
    from docx.shared import Cm  # type: ignore[import]

    left = _requisites_lines(exe, "Исполнитель")
    right = _requisites_lines(cli, "Заказчик")
    n_rows = max(len(left), len(right))
    left += [""] * (n_rows - len(left))
    right += [""] * (n_rows - len(right))

    tbl = doc.add_table(rows=n_rows, cols=2)
    tbl.autofit = False
    for i, (l, r) in enumerate(zip(left, right)):
        tbl.cell(i, 0).text = l
        tbl.cell(i, 1).text = r
    # Колонки 8.5 + 8.5 см.
    _set_column_widths(tbl, (8.5, 8.5))
    _remove_table_borders(tbl)
    # Первую строку (роль) выделяем жирным.
    for col in (0, 1):
        for run in tbl.cell(0, col).paragraphs[0].runs:
            run.bold = True


def _emit_signatures_block(doc, exe: Party, cli: Party) -> None:
    """Линия подписи и расшифровка: слева Исполнитель / справа Заказчик."""
    doc.add_paragraph("Подписи сторон:")
    doc.add_paragraph()
    tbl = doc.add_table(rows=3, cols=2)
    tbl.autofit = False

    tbl.cell(0, 0).text = (
        f"От Исполнителя: {exe.signatory_position or '___'}"
    )
    tbl.cell(0, 1).text = (
        f"От Заказчика:   {cli.signatory_position or '___'}"
    )
    tbl.cell(1, 0).text = "_________/_________________/"
    tbl.cell(1, 1).text = "_________/_________________/"
    tbl.cell(2, 0).text = f"     ({exe.signatory_name or 'ФИО'})"
    tbl.cell(2, 1).text = f"     ({cli.signatory_name or 'ФИО'})"

    _set_column_widths(tbl, (8.5, 8.5))
    _remove_table_borders(tbl)


def _remove_table_borders(tbl) -> None:
    from docx.oxml.ns import qn  # type: ignore[import]
    from docx.oxml import OxmlElement  # type: ignore[import]
    tbl_el = tbl._tbl
    tbl_pr = tbl_el.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tbl_pr)
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tbl_pr.append(borders)


def _set_column_widths(tbl, widths_cm: tuple[float, ...]) -> None:
    from docx.shared import Cm  # type: ignore[import]
    for col, w in enumerate(widths_cm):
        for cell in tbl.column_cells(col):
            cell.width = Cm(w)


# ───────────────────────────── MD builder ───────────────────────────────

def build_md(meta: ContractMeta, out_path: Path) -> Path:
    """Параллельная MD-версия — для diff/review без Word."""
    lines: list[str] = []
    lines.append(f"# {meta.title()}")
    lines.append("")
    lines.append("по сбору и систематизации данных для целей оценки")
    lines.append("")
    lines.append(f"**г. {meta.city or '___'}** — {meta.date_str}")
    lines.append("")
    lines.append(party_header_line(meta.executor, "Исполнитель") + ",")
    lines.append("с одной стороны, и")
    lines.append("")
    lines.append(party_header_line(meta.client, "Заказчик") + ",")
    lines.append("с другой стороны, далее совместно — «Стороны», заключили")
    lines.append("настоящий Договор о нижеследующем:")
    lines.append("")

    sections = [
        _section_predmet(meta.kind),
        _section_price(),
        _section_terms(),
        _section_rights(),
        _section_confid(),
        _section_liability(),
        _section_disputes(),
        _section_final(meta.kind),
        _section_appendices(),
    ]
    for heading, paragraphs in sections:
        lines.append(f"## {heading}")
        lines.append("")
        for p in paragraphs:
            lines.append(p)
            lines.append("")

    lines.append("## 9. РЕКВИЗИТЫ СТОРОН")
    lines.append("")
    lines.append("| Исполнитель | Заказчик |")
    lines.append("|---|---|")
    left = _requisites_lines(meta.executor, "Исполнитель")
    right = _requisites_lines(meta.client, "Заказчик")
    n = max(len(left), len(right))
    left += [""] * (n - len(left)); right += [""] * (n - len(right))
    for l, r in zip(left, right):
        lines.append(f"| {l} | {r} |")
    lines.append("")
    lines.append("**Подписи сторон:**")
    lines.append("")
    lines.append("| От Исполнителя | От Заказчика |")
    lines.append("|---|---|")
    lines.append(
        f"| {meta.executor.signatory_position or '___'} "
        f"| {meta.client.signatory_position or '___'} |"
    )
    lines.append("| _________/_________/ | _________/_________/ |")
    lines.append(
        f"| ({meta.executor.signatory_name or 'ФИО'}) "
        f"| ({meta.client.signatory_name or 'ФИО'}) |"
    )
    lines.append("")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(lines), encoding="utf-8")
    return out_path


# ───────────────────────────── main ─────────────────────────────────────

def main() -> int:
    cp(f"\n{BANNER}", C.B)
    cp("  ekcelo · Тело рамочного договора (DOCX + MD)", C.B)
    cp(f"{BANNER}\n", C.B)

    # 1) Выгрузка
    out_dir = ask_path(
        "Папка для сохранения DOCX/MD",
        must_exist=False, default=str(Path.cwd() / "out"),
    )
    out_dir.mkdir(parents=True, exist_ok=True)

    # 2) Опциональная папка проекта (только для поиска parties.json)
    proj_raw = ask("Папка проекта (Enter — пропустить)", "")
    project = Path(proj_raw).expanduser() if proj_raw else None
    if project and not project.is_dir():
        cp(f"  ⚠ Папка не найдена: {project} — игнорирую.", C.Y)
        project = None

    # 3) parties.json
    execs, clis, src = load_parties_file(project)
    if src:
        cp(f"\n· parties.json: {src}", C.CY)
        cp(f"  Исполнителей в файле: {len(execs)}; Заказчиков: {len(clis)}",
           C.CY)
    else:
        cp("\n· parties.json не найден — реквизиты можно ввести вручную "
           "или оставить «___».", C.Y)

    executor = pick_or_enter_party("Исполнитель", execs)
    client = pick_or_enter_party("Заказчик", clis)

    # 4) Номер / дата / город
    now = datetime.now()
    default_num = f"{now.strftime('%Y%m%d')}-1"
    number = ask("\nНомер договора", default=default_num)
    date_str = ask("Дата договора", default=ru_date(now))
    # Город по умолчанию — из адреса Исполнителя (грубая эвристика).
    default_city = _guess_city(executor.address) or "Ростов-на-Дону"
    city = ask("Город заключения", default=default_city)

    # 5) Тип договора
    kind_idx = ask_choice(
        "\nТип договора:",
        [
            "Возмездные услуги по ГК РФ (гл. 39, ст. 779-783)",
            "Бесшовный маппинг под отчёт по 135-ФЗ "
            "(основной — оценка по 135-ФЗ)",
        ],
        default=1,
    )
    kind = CONTRACT_KIND_GK if kind_idx == 1 else CONTRACT_KIND_FZ135

    meta = ContractMeta(
        number=number, date_str=date_str, city=city, kind=kind,
        executor=executor, client=client,
    )

    # 6) Генерация файлов
    ts = now.strftime("%Y%m%d_%H%M%S")
    safe_num = number.replace("/", "-").replace(" ", "_")
    out_md = out_dir / f"Contract_{safe_num}_{ts}.md"
    out_docx = out_dir / f"Contract_{safe_num}_{ts}.docx"

    build_md(meta, out_md)
    cp(f"\n✓ MD : {out_md}", C.G)

    try:
        build_docx(meta, out_docx)
        cp(f"✓ DOCX: {out_docx}", C.G)
    except ImportError as e:
        cp(f"\n⚠ python-docx не установлен: {e}", C.Y)
        cp("  pip install python-docx", C.Y)
        cp(f"  DOCX пропущен; MD доступен: {out_md}", C.Y)
    return 0


def _guess_city(address: str) -> str:
    """Грубо извлекает первую «г. Город» из адреса. Иначе ''."""
    if not address:
        return ""
    import re
    m = re.search(r"г\.\s*([А-ЯЁ][А-Яа-яЁё\-]+)", address)
    return m.group(1) if m else ""


if __name__ == "__main__":
    try:
        sys.exit(main())
    except KeyboardInterrupt:
        cp("\n✗ Прервано пользователем.", C.Y)
        sys.exit(130)
