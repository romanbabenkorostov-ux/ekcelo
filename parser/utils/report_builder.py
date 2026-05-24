"""ReportBuilder Protocol + два рендерера (MD / DOCX-native).

Реализует dev/SPEC_TEMPORAL_REPORTS.md §9 + §10 + §17:
  • `ReportBuilder` Protocol — единый интерфейс для отчётов 09_v1.
  • `MarkdownBuilder` — git-friendly MD-output.
  • `DocxNativeBuilder` — DOCX через python-docx с SEQ-полями (Рисунок N),
    TOC, COM-обновлением (заимствует из 06_photo_report_to_docx_v3.py
    §17.2/§17.3/§17.4).
  • `SourceTracker` — нумерация footnotes (§10.3).
"""
from __future__ import annotations

from pathlib import Path
from typing import Protocol


# ─── SourceTracker (§10.3) ──────────────────────────────────────────────────


class SourceTracker:
    """Накопитель источников: source_id → footnote-номер.

    `ref(source_id, description)` возвращает `[^N]`-маркер для inline-вставки;
    повторный вызов с тем же source_id даёт тот же N (не дублируется в блоке).
    `render_block()` — финальный `<details>`-блок со списком источников.
    """

    def __init__(self) -> None:
        self._map: dict[str, int] = {}
        self._descs: list[str] = []

    def ref(self, source_id: str, description: str) -> str:
        if source_id not in self._map:
            self._map[source_id] = len(self._descs) + 1
            self._descs.append(description)
        return f"[^{self._map[source_id]}]"

    def render_block(self) -> str:
        if not self._descs:
            return ""
        lines = ["<details>",
                 "<summary>Источники (служебный блок, "
                 "нумерация для внутрифайлового использования)</summary>",
                 ""]
        for i, desc in enumerate(self._descs, 1):
            lines.append(f"[^{i}]: {desc}")
        lines.append("")
        lines.append("</details>")
        return "\n".join(lines)


# ─── ReportBuilder Protocol (§9.1) ──────────────────────────────────────────


class ReportBuilder(Protocol):
    def heading(self, text: str, level: int) -> None: ...
    def paragraph(self, text: str) -> None: ...
    def table(self, headers: list[str], rows: list[list[str]],
              title: str | None = None) -> None: ...
    def sources_block(self) -> None: ...
    def save(self, out_path: Path) -> Path: ...


# ─── MarkdownBuilder ────────────────────────────────────────────────────────


class MarkdownBuilder:
    """Git-friendly MD-output. Соответствует §10.1 footnotes-схеме."""

    def __init__(self, tracker: SourceTracker | None = None,
                 title: str | None = None) -> None:
        self.tracker = tracker or SourceTracker()
        self._lines: list[str] = []
        if title:
            self._lines.append(f"# {title}")
            self._lines.append("")

    def heading(self, text: str, level: int) -> None:
        level = max(1, min(level, 6))
        self._lines.append("")
        self._lines.append(f"{'#' * level} {text}")
        self._lines.append("")

    def paragraph(self, text: str) -> None:
        self._lines.append(text)
        self._lines.append("")

    def table(self, headers: list[str], rows: list[list[str]],
              title: str | None = None) -> None:
        if title:
            self._lines.append(f"**{title}**")
            self._lines.append("")
        self._lines.append("| " + " | ".join(headers) + " |")
        self._lines.append("|" + "|".join(["---"] * len(headers)) + "|")
        for row in rows:
            cells = [str(c).replace("|", "\\|") for c in row]
            self._lines.append("| " + " | ".join(cells) + " |")
        self._lines.append("")

    def sources_block(self) -> None:
        block = self.tracker.render_block()
        if block:
            self._lines.append("")
            self._lines.append("---")
            self._lines.append("")
            self._lines.append(block)

    def save(self, out_path: Path) -> Path:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text("\n".join(self._lines), encoding="utf-8")
        return out_path


# ─── DocxNativeBuilder (§9.2 + §17 заимствования из 06) ─────────────────────


class DocxNativeBuilder:
    """DOCX-native через python-docx; SEQ для нумерации; TOC; COM-обновление.

    Заимствует паттерны из parser/scripts/pirushin_sosn_rocha_06_photo_report_to_docx_v3.py:
      • _insert_seq_field — §17.2 (Рисунок/Таблица N автонумерация).
      • TOC через OxmlElement — §17.3.
      • update_fields_via_word — §17.4 (Win10-only post-step).

    На Linux/CI всё работает без COM — поля обновятся при первом
    открытии в Word (F9 не обязателен — Word сам обновит, dirty=true).
    """

    def __init__(self, tracker: SourceTracker | None = None,
                 title: str | None = None) -> None:
        try:
            from docx import Document  # type: ignore[import]
            from docx.shared import Cm, Pt  # type: ignore[import]
        except ImportError as e:
            raise RuntimeError(
                "python-docx не установлен: pip install python-docx"
            ) from e

        self.tracker = tracker or SourceTracker()
        self._docx = Document()
        # A4 + поля 2 см (заимствовано из 06 setup_document)
        for section in self._docx.sections:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)
        style = self._docx.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(10)
        if title:
            self._docx.add_heading(title, level=1)
        self._table_seq_inserted = False

    # --- внутренние OxmlElement helpers (заимствовано из 06:320-352, 429-479) ---

    def _add_field(self, paragraph, instr: str, cached: str = "") -> None:
        from docx.oxml import OxmlElement  # type: ignore[import]
        from docx.oxml.ns import qn  # type: ignore[import]
        r = paragraph.add_run()
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        fld.set(qn("w:dirty"), "true")
        r._element.append(fld)

        r = paragraph.add_run()
        it = OxmlElement("w:instrText")
        it.set(qn("xml:space"), "preserve")
        it.text = instr
        r._element.append(it)

        r = paragraph.add_run()
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "separate")
        r._element.append(fld)

        if cached:
            paragraph.add_run(cached)

        r = paragraph.add_run()
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "end")
        r._element.append(fld)

    # --- ReportBuilder Protocol ---

    def heading(self, text: str, level: int) -> None:
        level = max(1, min(level, 9))
        self._docx.add_heading(text, level=level)

    def paragraph(self, text: str) -> None:
        self._docx.add_paragraph(text)

    def table(self, headers: list[str], rows: list[list[str]],
              title: str | None = None) -> None:
        if title:
            # «Таблица N. <title>» с SEQ-полем для автонумерации
            p = self._docx.add_paragraph()
            p.add_run("Таблица ")
            self._add_field(p, " SEQ Таблица \\* ARABIC ", cached="#")
            p.add_run(f". {title}")

        n_cols = len(headers)
        n_rows = len(rows) + 1
        tbl = self._docx.add_table(rows=n_rows, cols=n_cols)
        tbl.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            tbl.cell(0, i).text = str(h)
        for ri, row in enumerate(rows, 1):
            for ci, cell in enumerate(row):
                tbl.cell(ri, ci).text = str(cell)

    def sources_block(self) -> None:
        """Источники как развёрнутый параграф (DOCX не понимает `<details>`)."""
        if not self.tracker._descs:
            return
        self._docx.add_heading("Источники", level=2)
        self._docx.add_paragraph(
            "Служебный блок, нумерация для внутрифайлового использования."
        ).italic = True
        for i, desc in enumerate(self.tracker._descs, 1):
            self._docx.add_paragraph(f"[{i}] {desc}")

    def save(self, out_path: Path) -> Path:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        self._docx.save(str(out_path))
        # Win10 only: попытка обновить поля через Word COM (заимствовано из 06).
        try:
            self._update_fields_via_word(out_path)
        except Exception:
            pass  # graceful: на Linux/CI Word недоступен — поля обновятся при открытии
        return out_path

    def _update_fields_via_word(self, docx_path: Path) -> bool:
        """COM-обновление полей (Win10 + pywin32). Заимствовано из 06:575-646."""
        try:
            import win32com.client  # type: ignore[import]
            import pythoncom  # type: ignore[import]
        except ImportError:
            return False
        word = None
        doc = None
        try:
            pythoncom.CoInitialize()
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            doc = word.Documents.Open(str(docx_path.resolve()), ReadOnly=False)
            try:
                doc.Fields.Update()
            except Exception:
                pass
            try:
                for _ in range(2):
                    for toc in doc.TablesOfContents:
                        toc.Update()
            except Exception:
                pass
            doc.Save()
            return True
        finally:
            try:
                if doc is not None:
                    doc.Close(SaveChanges=False)
            except Exception:
                pass
            try:
                if word is not None:
                    word.Quit()
            except Exception:
                pass
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass
